import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document

import importlib.util


INBOX_DIR = Path("inbox")
PROCESSED_DIR = Path("processed") / "cblm"
STATE_DIR = Path("state") / "ia_payloads"
OUTPUT_DIR = Path("output") / "ia"


LO_HEAD_RE = re.compile(r"^LEARNING OUTCOME NO\.\s*(\d+)\s*-\s*(.+)$", re.I)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _safe_filename(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "_", (text or "").strip())
    text = re.sub(r"_+", "_", text).strip("_")
    return text or "Course"


def _list_inbox_unit_docx() -> list[Path]:
    if not INBOX_DIR.exists():
        return []
    files = [p for p in INBOX_DIR.glob("CBLM_Unit_*.docx") if p.is_file()]
    return sorted(files, key=lambda p: p.name.lower())


def _parse_filename(path: Path) -> tuple[int, str, str]:
    """
    Expected: CBLM_Unit_<n>_<COURSE_CODE>_<Title>.docx
    COURSE_CODE may contain underscores (e.g., ICC_5, IIPC_13).
    Infer COURSE_CODE from the first 1–2 tokens after the unit number:
    - Prefer <LETTERS>_<DIGITS> (e.g., ICC_5, IIPC_13, IS_6, IPE_10)
    - Otherwise accept <LETTERS><DIGITS> (e.g., ICC14)
    """
    m = re.match(r"^CBLM_Unit_(\d+)_(.+)\.docx$", path.name, re.I)
    if not m:
        raise ValueError(f"Invalid filename (expected CBLM_Unit_<n>_*.docx): {path.name}")
    unit_index = int(m.group(1))
    remainder = m.group(2)
    tokens = remainder.split("_")
    if len(tokens) < 3:
        raise ValueError(f"Invalid filename (expected CBLM_Unit_<n>_<COURSE_CODE>_<Title>.docx): {path.name}")

    if len(tokens) >= 2 and re.match(r"^[A-Za-z]+_[0-9]+$", f"{tokens[0]}_{tokens[1]}"):
        course_code = f"{tokens[0]}_{tokens[1]}"
        title_tokens = tokens[2:]
    elif re.match(r"^[A-Za-z]+[0-9]+$", tokens[0]):
        course_code = tokens[0]
        title_tokens = tokens[1:]
    else:
        course_code = tokens[0]
        title_tokens = tokens[1:]

    title = "_".join(title_tokens)
    return unit_index, course_code, title


def _extract_front_matter_from_tables(doc: Document) -> dict:
    out = {"qualification_title": "", "unit_of_competency": ""}

    def cell_text(cell) -> str:
        return _norm(getattr(cell, "text", "") or "")

    labelish = {
        "sector",
        "qualification title",
        "qualification",
        "unit of competency",
        "module title",
        "module descriptor",
        "prepared by",
        "code",
        "no.",
    }

    def is_label(value: str) -> bool:
        low = _norm(value).lower()
        if not low:
            return False
        if low in ("code", "module title", "unit of competency", "qualification title", "sector"):
            return True
        return any(k in low for k in labelish)

    def pick_value(texts: list[str], start_idx: int) -> str:
        for cand in texts[start_idx + 1 :]:
            cand = _norm(cand)
            if not cand:
                continue
            if is_label(cand):
                continue
            return cand
        for cand in reversed(texts):
            cand = _norm(cand)
            if cand and not is_label(cand):
                return cand
        return ""

    for table in doc.tables:
        for row in table.rows:
            texts = [cell_text(c) for c in row.cells]
            for idx, t in enumerate(texts):
                low = t.lower()
                if "qualification title" in low and idx + 1 < len(texts):
                    out["qualification_title"] = out["qualification_title"] or pick_value(texts, idx)
                if "unit of competency" in low and idx + 1 < len(texts):
                    out["unit_of_competency"] = out["unit_of_competency"] or pick_value(texts, idx)
    return out


def _extract_learning_outcomes(doc: Document) -> list[dict]:
    los: list[dict] = []
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        t = _norm(paragraphs[i].text)
        m = LO_HEAD_RE.match(t)
        if not m:
            i += 1
            continue

        lo = {"index": len(los) + 1, "title": _norm(m.group(2)), "contents": []}

        j = i + 1
        while j < len(paragraphs) and not _norm(paragraphs[j].text).lower().startswith("contents"):
            j += 1

        if j < len(paragraphs) and _norm(paragraphs[j].text).lower().startswith("contents"):
            k = j + 1
            while k < len(paragraphs):
                pk = paragraphs[k]
                tk = _norm(pk.text)
                if not tk:
                    k += 1
                    continue
                if LO_HEAD_RE.match(tk):
                    break
                if tk.lower().startswith("assessment criteria"):
                    break

                style = (getattr(pk.style, "name", "") or "").lower()
                if style.startswith("list"):
                    lo["contents"].append({"index": len(lo["contents"]) + 1, "title": tk})
                    k += 1
                    continue
                break

        los.append(lo)
        i = j + 1

    return los


def _build_course_ia_payload(units: list[tuple[int, str, str, Path]]) -> dict:
    # Determine qualification title (best effort): prefer first unit's table-derived qualification title,
    # fallback to filename title.
    qualification_title = ""

    aggregated_los: list[dict] = []
    lo_index = 1

    for unit_index, _course_code, title, path in units:
        doc = Document(str(path))

        fm = _extract_front_matter_from_tables(doc)
        if not qualification_title:
            qualification_title = fm.get("qualification_title") or title.replace("_", " ").strip()

        unit_of_competency = fm.get("unit_of_competency") or ""
        unit_of_competency = _norm(unit_of_competency)

        for lo in _extract_learning_outcomes(doc):
            lo_title = _norm(lo.get("title", ""))
            merged_title = f"{unit_of_competency} — {lo_title}" if unit_of_competency else lo_title
            aggregated_los.append({"index": lo_index, "title": merged_title, "contents": lo.get("contents", []) or []})
            lo_index += 1

    qualification_title = _norm(qualification_title) or "Qualification"

    return {
        "qualification_title": qualification_title,
        "current_unit": {
            "unit_of_competency": "",
            "module_title": "",
            "next_unit_of_competency": "",
            "learning_outcomes": aggregated_los,
            "ia": {},
        },
    }


def _outlines_for_units(units: list[tuple[int, str, str, Path]]) -> list[str]:
    outlines: list[str] = []
    for unit_index, _course_code, title, path in units:
        doc = Document(str(path))
        fm = _extract_front_matter_from_tables(doc)
        uc = _norm(fm.get("unit_of_competency") or "")
        qual = _norm(fm.get("qualification_title") or title.replace("_", " ").strip())
        parts = []
        parts.append(f"Qualification Title: {qual}")
        parts.append(f"Unit of Competency: {uc}")
        for lo in _extract_learning_outcomes(doc):
            lo_title = _norm(lo.get("title", ""))
            if not lo_title:
                continue
            parts.append(f"- Topic: {lo_title}")
            subtopics = [ _norm(c.get("title","")) for c in (lo.get("contents") or []) if _norm(c.get("title","")) ]
            if subtopics:
                parts.append("  Subtopics: " + "; ".join(subtopics))
        outlines.append("\n".join(parts).strip())
    return outlines


def _payloads_for_exam_generation(units: list[tuple[int, str, str, Path]]) -> list[dict]:
    """
    Build minimal payload-like dicts per unit so the local exam generator can extract UC/Topic/Subtopic.
    """
    out: list[dict] = []
    for unit_index, _course_code, title, path in units:
        doc = Document(str(path))
        fm = _extract_front_matter_from_tables(doc)
        uc = _norm(fm.get("unit_of_competency") or "")
        los = _extract_learning_outcomes(doc)
        out.append(
            {
                "current_unit": {
                    "index": unit_index,
                    "unit_of_competency": uc,
                    "learning_outcomes": los,
                }
            }
        )
    return out


def _load_generate_midterm_and_finals_local():
    tool_path = Path(__file__).resolve().parent / "generate_course_exams_local.py"
    spec = importlib.util.spec_from_file_location("generate_course_exams_local", tool_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load local exam generator module from: {tool_path}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    fn = getattr(mod, "generate_midterm_and_finals_local", None)
    if not callable(fn):
        raise RuntimeError("generate_midterm_and_finals_local not found in generate_course_exams_local.py")
    return fn




def main(argv: list[str]) -> int:
    unit_files = _list_inbox_unit_docx()
    if not unit_files:
        return 0

    parsed = []
    for p in unit_files:
        try:
            unit_index, course_code, title = _parse_filename(p)
        except Exception:
            # Skip invalid names; the prompt expects strict naming.
            continue
        parsed.append((unit_index, course_code, title, p))

    if not parsed:
        return 0

    course_codes = sorted({course_code for _, course_code, _, _ in parsed})
    chosen_course = course_codes[0]
    chosen_units = [(unit_index, course_code, title, p) for unit_index, course_code, title, p in parsed if course_code == chosen_course]
    chosen_units.sort(key=lambda x: x[0])
    if not chosen_units:
        return 0

    out_dir = OUTPUT_DIR / chosen_course
    state_dir = STATE_DIR / chosen_course
    processed_dir = PROCESSED_DIR / chosen_course
    out_dir.mkdir(parents=True, exist_ok=True)
    state_dir.mkdir(parents=True, exist_ok=True)
    processed_dir.mkdir(parents=True, exist_ok=True)

    payload = _build_course_ia_payload(chosen_units)

    # Exams are required for IA assembly; generate MIDTERM + FINAL locally (no external API).
    # This IA-only path uses the extracted LO/contents lists as the basis for course IA payload.
    generate_midterm_and_finals_local = _load_generate_midterm_and_finals_local()
    payload["exams"] = generate_midterm_and_finals_local(
        _payloads_for_exam_generation(chosen_units),
        course_code=chosen_course,
        course_title=payload.get("qualification_title", chosen_course),
        seed=0,
    )
    payload_path = state_dir / "IA_FULL.json"
    out_path = out_dir / "IA_FULL.docx"
    payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    result = subprocess.run(
        [str(Path(".venv") / "Scripts" / "python.exe"), str(Path("tools") / "assemble_ia.py"), str(payload_path), str(out_path)],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        sys.stderr.write(result.stdout)
        sys.stderr.write(result.stderr)
        return result.returncode

    moved: list[Path] = []
    for _, _, _, p in chosen_units:
        dest = processed_dir / p.name
        if dest.exists():
            dest = processed_dir / f"{p.stem}__DUPLICATE{p.suffix}"
        shutil.move(str(p), str(dest))
        moved.append(dest)

    summary = {
        "course_code": chosen_course,
        "units_processed": [p.name for p in moved],
        "ia_payload": str(payload_path),
        "ia_full_path": str(out_path),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
