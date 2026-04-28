import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docxcompose.composer import Composer


FONT_NAME = "Bookman Old Style"
FONT_SIZE = 12
TEMPLATE_FILES = {
    "front_matter": "00_front_matter.docx",
    "lo_intro": "10_lo_intro.docx",
    "learning_experience": "20_learning_experience.docx",
    "key_facts": "30_key_facts.docx",
    "lets_exercise": "40_lets_exercise.docx",
    "lets_apply": "50_lets_apply.docx",
}


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    return str(value)


def set_run_font(run):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)


def copy_run_style(source_run, target_run):
    if source_run is None:
        set_run_font(target_run)
        return
    set_run_font(target_run)
    try:
        target_run.style = source_run.style
    except Exception:
        pass
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name or FONT_NAME
    if target_run._element.rPr is not None:
        target_run._element.rPr.rFonts.set(qn("w:eastAsia"), source_run.font.name or FONT_NAME)
    target_run.font.size = source_run.font.size or Pt(FONT_SIZE)


def enforce_document_font(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = FONT_NAME
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                style.font.size = Pt(FONT_SIZE)
            except Exception:
                pass


def clear_paragraph(paragraph):
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)


def write_text_to_paragraph(paragraph, text):
    style_source_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph(paragraph)
    lines = safe_text(text).split("\n")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        copy_run_style(style_source_run, run)
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def set_cell_text(cell, text):
    first_paragraph = cell.paragraphs[0]
    write_text_to_paragraph(first_paragraph, text)
    for paragraph in cell.paragraphs[1:]:
        delete_paragraph(paragraph)


def replace_in_paragraphs(paragraphs, placeholder, value):
    replacement = safe_text(value)
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            write_text_to_paragraph(paragraph, paragraph.text.replace(placeholder, replacement))


def replace_in_table(table, placeholder, value):
    for row in table.rows:
        for cell in row.cells:
            replace_everywhere(cell, placeholder, value)


def replace_everywhere(doc_part, placeholder, value):
    replace_in_paragraphs(doc_part.paragraphs, placeholder, value)
    for table in doc_part.tables:
        replace_in_table(table, placeholder, value)


def apply_scalar_map(doc, values):
    for key, value in values.items():
        replace_everywhere(doc, "{{" + key + "}}", value)


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_row(table, row_index):
    tbl = table._tbl
    tr = table.rows[row_index]._tr
    tbl.remove(tr)


def normalize_all_runs_font(doc_part):
    for paragraph in doc_part.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_all_runs_font(cell)


def split_paragraph_blocks(text):
    text = safe_text(text).strip()
    if not text:
        return []
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    return blocks or [text]


def insert_paragraph_after_preserve(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def append_centered_image(doc, image_path, caption=None, width_inches=5.8):
    image_path = Path(image_path)
    if not image_path.exists():
        print(f"Warning: image not found, skipping: {image_path}", file=sys.stderr)
        return

    spacer = doc.add_paragraph("")
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    set_run_font(run)

    if caption and str(caption).strip():
        caption_paragraph = doc.add_paragraph(str(caption).strip())
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in caption_paragraph.runs:
            r.italic = True
            set_run_font(r)


def insert_centered_image_after(paragraph, image_path, caption=None, width_inches=5.8):
    image_path = Path(image_path)
    if not image_path.exists():
        print(f"Warning: image not found, skipping: {image_path}", file=sys.stderr)
        return paragraph

    img_paragraph = insert_paragraph_after_preserve(paragraph)
    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    set_run_font(run)

    last = img_paragraph
    if caption and str(caption).strip():
        cap_paragraph = insert_paragraph_after_preserve(img_paragraph)
        cap_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_paragraph.add_run(str(caption).strip())
        run.italic = True
        set_run_font(run)
        last = cap_paragraph

    return last


def replace_keyfacts_with_inline_image(doc_part, placeholder, key_facts_text, image_path=None, caption=None, width_inches=5.8):
    blocks = split_paragraph_blocks(key_facts_text)
    if not blocks:
        blocks = [""]

    replaced_any = False
    for paragraph in list(doc_part.paragraphs):
        if placeholder not in paragraph.text:
            continue

        replaced_any = True
        # Write first block into the existing paragraph (preserves layout position).
        first = paragraph.text.replace(placeholder, blocks[0])
        write_text_to_paragraph(paragraph, first)

        tail_anchor = paragraph
        # Insert image after the first block so it sits "inline" within Key Facts flow.
        if image_path:
            tail_anchor = insert_centered_image_after(tail_anchor, image_path, caption=caption, width_inches=width_inches)

        # Insert remaining blocks as new paragraphs after the image (or after first paragraph).
        for block in blocks[1:]:
            next_paragraph = insert_paragraph_after_preserve(tail_anchor)
            # Carry the same paragraph style for a consistent look.
            try:
                next_paragraph.style = paragraph.style
            except Exception:
                pass
            write_text_to_paragraph(next_paragraph, block)
            tail_anchor = next_paragraph

    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                if replace_keyfacts_with_inline_image(cell, placeholder, key_facts_text, image_path=image_path, caption=caption, width_inches=width_inches):
                    replaced_any = True

    return replaced_any


def load_payload(path):
    with Path(path).open("r", encoding="utf-8") as handle:
        return json.load(handle)


def validate_payload(payload):
    required_root = ["sector", "qualification_title", "qualification_units", "current_unit"]
    for key in required_root:
        if key not in payload:
            raise ValueError(f"Missing root field: {key}")

    unit = payload["current_unit"]
    required_unit = [
        "index",
        "unit_of_competency",
        "module_title",
        "next_unit_of_competency",
        "Module_Descriptor",
        "Laboratory",
        "training_materials",
        "learning_outcomes",
    ]
    for key in required_unit:
        if key not in unit:
            raise ValueError(f"Missing current_unit field: {key}")

    descriptor_words = len(safe_text(unit["Module_Descriptor"]).split())
    if not 80 <= descriptor_words <= 120:
        raise ValueError(f"Module_Descriptor must be 80-120 words, got {descriptor_words}")

    if not payload["qualification_units"]:
        raise ValueError("qualification_units must not be empty")
    if not unit["learning_outcomes"]:
        raise ValueError("current_unit.learning_outcomes must not be empty")

    for lo in unit["learning_outcomes"]:
        if "title" not in lo or "contents" not in lo:
            raise ValueError("Each learning outcome needs title and contents")
        if not lo["contents"]:
            raise ValueError(f"Learning outcome '{lo.get('title', '')}' has no contents")
        for content in lo["contents"]:
            for field in [
                "title",
                "key_facts",
                "exercise_questions",
                "answer_key",
                "apply_title",
                "apply_objective",
                "apply_sup_mat",
                "apply_equipment_list",
                "apply_steps_list",
                "apply_assessmentmethod",
                "apply_pc1",
                "apply_pc2",
                "apply_pc3",
                "apply_pc4",
                "apply_pc5",
            ]:
                if not safe_text(content.get(field, "")).strip():
                    raise ValueError(f"Missing content field '{field}' in LO '{lo['title']}'")
            if len(safe_text(content["key_facts"]).split()) < 600:
                raise ValueError(f"Key Facts too short for content '{content['title']}'")
            answer_lines = [line for line in safe_text(content["answer_key"]).splitlines() if line.strip()]
            if len(answer_lines) != 10:
                raise ValueError(f"Expected 10 answer key lines for content '{content['title']}'")


def get_front_matter_learning_outcome_items(unit):
    items = []
    for lo in unit["learning_outcomes"]:
        contents = lo.get("contents") or []
        if contents:
            items.extend(safe_text(content.get("title", "")).strip() for content in contents if safe_text(content.get("title", "")).strip())
        else:
            title = safe_text(lo.get("title", "")).strip()
            if title:
                items.append(title)
    return items


def render_front_matter(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)

    unit = payload["current_unit"]
    apply_scalar_map(
        doc,
        {
            "sector": payload["sector"],
            "qualification_title": payload["qualification_title"],
            "unit_of_competency_x": unit["unit_of_competency"],
            "module_title": unit["module_title"],
            "next_unit_of_competency": unit["next_unit_of_competency"],
            "Module_Descriptor": unit["Module_Descriptor"],
        },
    )

    if len(doc.tables) >= 2:
        table = doc.tables[1]
        units = payload["qualification_units"]
        row_index = 1
        while row_index < len(table.rows):
            if row_index <= len(units):
                item = units[row_index - 1]
                row = table.rows[row_index]
                set_cell_text(row.cells[0], str(item.get("index", row_index)))
                set_cell_text(row.cells[1], safe_text(item["unit_of_competency"]))
                set_cell_text(row.cells[2], safe_text(item["module_title"]))
                set_cell_text(row.cells[3], safe_text(item["unit_of_competency_code"]))
            else:
                delete_row(table, row_index)
                continue
            row_index += 1

    lo_paragraphs = [p for p in doc.paragraphs if p.text.strip() == "{{LO_X}}"]
    learning_outcomes = get_front_matter_learning_outcome_items(unit)
    for index, paragraph in enumerate(lo_paragraphs):
        if index < len(learning_outcomes):
            write_text_to_paragraph(paragraph, learning_outcomes[index])
        else:
            delete_paragraph(paragraph)

    normalize_all_runs_font(doc)
    return doc


def render_lo_intro(template_path, payload, lo):
    doc = Document(template_path)
    enforce_document_font(doc)
    apply_scalar_map(
        doc,
        {
            "Y": lo["index"],
            "LO": lo["title"],
            "Laboratory": payload["current_unit"]["Laboratory"],
            "training_materials": payload["current_unit"]["training_materials"],
        },
    )

    content_paragraphs = [p for p in doc.paragraphs if p.text.strip() == "{{Contents_Z}}"]
    for index, paragraph in enumerate(content_paragraphs):
        if index < len(lo["contents"]):
            write_text_to_paragraph(paragraph, lo["contents"][index]["title"])
        else:
            delete_paragraph(paragraph)

    normalize_all_runs_font(doc)
    return doc


def render_learning_experience(template_path, unit_index, lo):
    doc = Document(template_path)
    enforce_document_font(doc)
    apply_scalar_map(doc, {"X": unit_index, "Y": lo["index"], "LO": lo["title"]})

    table = doc.tables[0]
    contents = lo["contents"]
    needed_rows = 1 + (len(contents) * 3)

    row_index = 1
    for content in contents:
        replacements = [
            ("Reading Key Facts {{X}}.{{Y}}-{{Z}}. {{Contents_Z}}", f"Reading Key Facts {unit_index}.{lo['index']}-{content['index']}. {content['title']}"),
            (
                "Answering Let’s Exercise {{X}}.{{Y}}-{{Z}}. {{Contents_Z}}\nCompare answers to Answer Key {{X}}.{{Y}}-{{Z}} {{Contents_Z}}",
                f"Answering Let’s Exercise {unit_index}.{lo['index']}-{content['index']}. {content['title']}\nCompare answers to Answer Key {unit_index}.{lo['index']}-{content['index']} {content['title']}",
            ),
            (
                "Performing Let’s Apply {{X}}.{{Y}}-{{Z}}.  {{la_title_z}}",
                f"Performing Let’s Apply {unit_index}.{lo['index']}-{content['index']}.  {content['apply_title']}",
            ),
        ]
        for replacement in replacements:
            set_cell_text(table.rows[row_index].cells[0], replacement[1])
            row_index += 1

    while len(table.rows) > needed_rows:
        delete_row(table, len(table.rows) - 1)

    normalize_all_runs_font(doc)
    return doc


def render_key_facts(template_path, unit_index, lo_index, content, payload_base_dir=None):
    doc = Document(template_path)
    enforce_document_font(doc)
    apply_scalar_map(
        doc,
        {
            "X": unit_index,
            "Y": lo_index,
            "Z": content["index"],
            "Contents": content["title"],
        },
    )

    payload_base_dir = Path(payload_base_dir or ".")
    image_path = content.get("key_facts_image_path") or content.get("image_path")
    caption = content.get("key_facts_image_caption") or content.get("image_caption")
    width_inches = content.get("key_facts_image_width_in") or content.get("image_width_in") or 5.8
    try:
        width_inches = float(width_inches)
    except Exception:
        width_inches = 5.8

    resolved = None
    if image_path:
        resolved = Path(image_path)
        if not resolved.is_absolute():
            resolved = payload_base_dir / resolved

    # Replace Key Facts placeholder and optionally insert image inline in the Key Facts flow.
    placeholder = "{{Contents_Key_Facts}}"
    replaced = replace_keyfacts_with_inline_image(
        doc,
        placeholder,
        content.get("key_facts", ""),
        image_path=resolved,
        caption=caption,
        width_inches=width_inches,
    )
    if not replaced:
        # Fallback: if the template doesn't contain the placeholder (unexpected), append at end.
        replace_everywhere(doc, placeholder, content.get("key_facts", ""))
        if resolved:
            append_centered_image(doc, resolved, caption=caption, width_inches=width_inches)

    normalize_all_runs_font(doc)
    return doc


def render_lets_exercise(template_path, unit_index, lo_index, content):
    doc = Document(template_path)
    enforce_document_font(doc)
    apply_scalar_map(
        doc,
        {
            "X": unit_index,
            "Y": lo_index,
            "Z": content["index"],
            "Contents_Z_LE_MC": content["exercise_questions"],
            "Contents_Z_LE_MC_Answer": content["answer_key"],
        },
    )
    normalize_all_runs_font(doc)
    return doc


def render_lets_apply(template_path, unit_index, lo_index, content):
    doc = Document(template_path)
    enforce_document_font(doc)
    apply_scalar_map(
        doc,
        {
            "X": unit_index,
            "Y": lo_index,
            "Z": content["index"],
            "la_title_z": content["apply_title"],
            "la_z_objective": content["apply_objective"],
            "la_z_sup_mat": content["apply_sup_mat"],
            "la_z_equipment_list": content["apply_equipment_list"],
            "la_z_steps_list": content["apply_steps_list"],
            "la_z_assessmentmethod": content["apply_assessmentmethod"],
            "la_z_pc1": content["apply_pc1"],
            "la_z_pc2": content["apply_pc2"],
            "la_z_pc3": content["apply_pc3"],
            "la_z_pc4": content["apply_pc4"],
            "la_z_pc5": content["apply_pc5"],
        },
    )
    normalize_all_runs_font(doc)
    return doc


def save_temp_doc(doc, directory, name):
    path = Path(directory) / name
    doc.save(path)
    return path


def append_with_page_break(composer, document_path):
    composer.doc.add_page_break()
    composer.append(Document(document_path))


def assemble_cblm(payload, templates_dir, output_path, payload_base_dir=None):
    unit = payload["current_unit"]
    templates_dir = Path(templates_dir)
    output_path = Path(output_path)
    payload_base_dir = Path(payload_base_dir or ".")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as temp_dir:
        front_doc = render_front_matter(templates_dir / TEMPLATE_FILES["front_matter"], payload)
        front_path = save_temp_doc(front_doc, temp_dir, "00_front.docx")
        composer = Composer(Document(front_path))

        part_index = 1
        for lo in unit["learning_outcomes"]:
            lo_doc = render_lo_intro(templates_dir / TEMPLATE_FILES["lo_intro"], payload, lo)
            lo_path = save_temp_doc(lo_doc, temp_dir, f"{part_index:03d}_lo_intro.docx")
            append_with_page_break(composer, lo_path)
            part_index += 1

            experience_doc = render_learning_experience(templates_dir / TEMPLATE_FILES["learning_experience"], unit["index"], lo)
            experience_path = save_temp_doc(experience_doc, temp_dir, f"{part_index:03d}_learning_experience.docx")
            append_with_page_break(composer, experience_path)
            part_index += 1

            for content in lo["contents"]:
                key_doc = render_key_facts(
                    templates_dir / TEMPLATE_FILES["key_facts"],
                    unit["index"],
                    lo["index"],
                    content,
                    payload_base_dir=payload_base_dir,
                )
                key_path = save_temp_doc(key_doc, temp_dir, f"{part_index:03d}_key_facts.docx")
                append_with_page_break(composer, key_path)
                part_index += 1

                exercise_doc = render_lets_exercise(templates_dir / TEMPLATE_FILES["lets_exercise"], unit["index"], lo["index"], content)
                exercise_path = save_temp_doc(exercise_doc, temp_dir, f"{part_index:03d}_exercise.docx")
                append_with_page_break(composer, exercise_path)
                part_index += 1

                apply_doc = render_lets_apply(templates_dir / TEMPLATE_FILES["lets_apply"], unit["index"], lo["index"], content)
                apply_path = save_temp_doc(apply_doc, temp_dir, f"{part_index:03d}_apply.docx")
                append_with_page_break(composer, apply_path)
                part_index += 1

        composer.save(output_path)


def main():
    if len(sys.argv) not in {3, 4}:
        print("Usage: python assemble_cblm.py <payload.json> <output.docx> [templates_dir]")
        return 1

    payload_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    templates_dir = Path(sys.argv[3]) if len(sys.argv) == 4 else Path("templates")

    payload = load_payload(payload_path)
    validate_payload(payload)
    assemble_cblm(payload, templates_dir, output_path, payload_base_dir=payload_path.parent)
    print(str(output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
