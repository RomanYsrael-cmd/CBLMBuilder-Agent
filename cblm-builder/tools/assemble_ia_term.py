import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docxcompose.composer import Composer


FONT_NAME = "Arial Narrow"
FONT_SIZE = 12

IA_TEMPLATE_FILES = {
    "front_page": "00_FrontPage.docx",
    "ia_plan": "01_iaplan.docx",
    "specific_instructions": "02_iaspesificinstruction.docx",
    "midterm": "03_midterm.docx",
    "finals": "04_finals.docx",
    "oral_questions": "05_iaquestionwithanswer.docx",
}


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    return str(value)


def set_run_font(run):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)


def enforce_document_font(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = FONT_NAME
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                style.font.size = Pt(FONT_SIZE)
            except Exception:
                pass


def clear_paragraph(paragraph):
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)


def write_text_to_paragraph(paragraph, text):
    style_source_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph(paragraph)
    lines = safe_text(text).split("\n")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run)
        if style_source_run is not None:
            run.bold = style_source_run.bold
            run.italic = style_source_run.italic
            run.underline = style_source_run.underline
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def replace_in_paragraphs(paragraphs, placeholder, value):
    replacement = safe_text(value)
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            write_text_to_paragraph(paragraph, paragraph.text.replace(placeholder, replacement))


def replace_everywhere(doc_part, placeholder, value):
    replace_in_paragraphs(doc_part.paragraphs, placeholder, value)
    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_everywhere(cell, placeholder, value)


def apply_scalar_map(doc, values):
    for key, value in values.items():
        replace_everywhere(doc, "{{" + key + "}}", value)


def delete_row(table, row_index):
    tbl = table._tbl
    tr = table.rows[row_index]._tr
    tbl.remove(tr)


def normalize_all_runs_font(doc_part):
    for paragraph in doc_part.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_all_runs_font(cell)


_COMMON_VERBS = {
    "apply",
    "analyze",
    "assess",
    "calculate",
    "communicate",
    "conduct",
    "configure",
    "coordinate",
    "create",
    "demonstrate",
    "describe",
    "design",
    "develop",
    "evaluate",
    "explain",
    "identify",
    "implement",
    "inspect",
    "install",
    "interpret",
    "maintain",
    "manage",
    "monitor",
    "operate",
    "perform",
    "plan",
    "prepare",
    "present",
    "record",
    "report",
    "test",
    "troubleshoot",
    "use",
}


def to_ability_phrase(content_title):
    title = safe_text(content_title).strip()
    if not title:
        return ""
    lower = title.lower()
    if lower.startswith("to "):
        title = title[3:].lstrip()
        lower = title.lower()
    first_word = lower.split()[0] if lower.split() else ""
    if first_word in _COMMON_VERBS:
        return title[:1].lower() + title[1:]
    return f"demonstrate competence in {title}"


def _get_unit(payload):
    if "current_unit" not in payload:
        raise ValueError("Payload missing required field: current_unit")
    return payload["current_unit"]


def _get_term(payload):
    term = safe_text(payload.get("term", "")).strip().upper()
    if term not in {"MIDTERM", "FINALS"}:
        raise ValueError("Payload missing valid top-level field: term")
    return term


def _get_oral_questions(payload):
    unit = _get_unit(payload)
    ia_block = unit.get("ia") or {}
    oral_questions = ia_block.get("oral_questions")
    if not isinstance(oral_questions, list):
        raise ValueError("Payload missing current_unit.ia.oral_questions")

    normalized = []
    for item in oral_questions:
        if not isinstance(item, dict):
            continue
        question = safe_text(item.get("question", "")).strip()
        acceptable_answer = safe_text(item.get("acceptable_answer", "")).strip()
        if question and acceptable_answer:
            normalized.append({"question": question, "acceptable_answer": acceptable_answer})
        if len(normalized) >= 5:
            break
    if len(normalized) != 5:
        raise ValueError("Payload must provide exactly 5 oral questions with acceptable answers")
    return normalized


def render_front_page(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    term = _get_term(payload)
    apply_scalar_map(
        doc,
        {
            "qualification": payload.get("qualification_title", ""),
            "qualification_title": payload.get("qualification_title", ""),
            "term": term,
            "TERM": term,
        },
    )
    normalize_all_runs_font(doc)
    return doc


def render_specific_instructions(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    term = _get_term(payload)
    apply_scalar_map(
        doc,
        {
            "qualification": payload.get("qualification_title", ""),
            "qualification_title": payload.get("qualification_title", ""),
            "term": term,
            "TERM": term,
            "page_count": payload.get("page_count", ""),
        },
    )
    normalize_all_runs_font(doc)
    return doc


def render_oral_questions(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)

    oral_questions = _get_oral_questions(payload)
    scalar_map = {}
    for idx, item in enumerate(oral_questions, start=1):
        scalar_map[f"oral_question_{idx}"] = item["question"]
        scalar_map[f"oral_question_{idx}_acceptable_answer"] = item["acceptable_answer"]

    apply_scalar_map(doc, scalar_map)
    normalize_all_runs_font(doc)
    return doc


def _subtopics_text(content):
    out = []
    for idx, sub in enumerate(content.get("subtopics") or [], start=1):
        if isinstance(sub, dict):
            title = safe_text(sub.get("title", "")).strip()
        else:
            title = safe_text(sub).strip()
        if not title:
            continue
        out.append(title)
    return "\n".join(out)


def _fill_plan_or_rating_table(table, unit):
    def row_has_any(row, needles):
        return any(any(n in cell.text for n in needles) for cell in row.cells)

    placeholders = ["{{LO}}", "{{Contents}}", "{{Contents_Z}}", "{{Y}}", "{{Z}}"]
    dynamic_row_indexes = [i for i, r in enumerate(table.rows) if row_has_any(r, placeholders)]
    dynamic_row_indexes = [
        i
        for i in dynamic_row_indexes
        if "{{unit_of_competency}}" not in "|".join(c.text for c in table.rows[i].cells)
        and "{{module_title}}" not in "|".join(c.text for c in table.rows[i].cells)
    ]

    los = unit.get("learning_outcomes", []) or []
    lo_iter = iter(los)
    current_lo = None
    current_contents_iter = iter(())
    used_row_indexes = []

    for row_index in dynamic_row_indexes:
        row = table.rows[row_index]
        row_text = "|".join(cell.text for cell in row.cells)

        if "{{LO}}" in row_text:
            current_lo = next(lo_iter, None)
            current_contents_iter = iter((current_lo.get("contents", []) or []) if current_lo else ())
            if not current_lo:
                break

            y = current_lo.get("index", "")
            lo_title = current_lo.get("title", "")
            for cell in row.cells:
                text = cell.text.replace("{{Y}}", safe_text(y)).replace("{{LO}}", safe_text(lo_title))
                write_text_to_paragraph(cell.paragraphs[0], text)
                for p in cell.paragraphs[1:]:
                    p._element.getparent().remove(p._element)
            used_row_indexes.append(row_index)
            continue

        if not current_lo:
            break

        content = next(current_contents_iter, None)
        if not content:
            continue

        y = current_lo.get("index", "")
        z = content.get("index", "")
        content_title = safe_text(content.get("title", ""))
        content_for_contents = to_ability_phrase(content_title)
        contents_list = _subtopics_text(content)

        for cell in row.cells:
            text = cell.text.replace("{{Y}}", safe_text(y)).replace("{{Z}}", safe_text(z))
            text = text.replace("{{Contents}}", content_for_contents)
            text = text.replace("{{Contents_Z}}", contents_list)
            write_text_to_paragraph(cell.paragraphs[0], text)
            for p in cell.paragraphs[1:]:
                p._element.getparent().remove(p._element)

        used_row_indexes.append(row_index)

    used = set(used_row_indexes)
    for row_index in reversed(dynamic_row_indexes):
        if row_index not in used:
            delete_row(table, row_index)


def render_ia_plan(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)
    term = _get_term(payload)

    apply_scalar_map(
        doc,
        {
            "qualification": payload.get("qualification_title", ""),
            "qualification_title": payload.get("qualification_title", ""),
            "term": term,
            "TERM": term,
        },
    )

    if doc.tables:
        _fill_plan_or_rating_table(doc.tables[0], unit)

    normalize_all_runs_font(doc)
    return doc


def render_exam(template_path, *, course_code: str, course_title: str, term: str, mcqs: list[dict]) -> Document:
    doc = Document(template_path)
    enforce_document_font(doc)
    apply_scalar_map(doc, {"COURSE_CODE": course_code, "COURSE_TITLE": course_title, "TERM": term, "term": term})

    question_table = None
    for table in doc.tables:
        if len(table.rows) == 25 and len(table.columns) == 2:
            question_table = table
            break
    if question_table is None:
        raise ValueError("Could not locate question table (expected 25x2).")

    cells = []
    for row in question_table.rows:
        cells.append(row.cells[0])
        cells.append(row.cells[1])

    if len(mcqs) != 50:
        raise ValueError(f"Expected exactly 50 MCQs, got {len(mcqs)}")

    for idx, item in enumerate(mcqs, start=1):
        mapping = {
            "Q_NUM": str(idx),
            "QUESTION": safe_text(item.get("question", "")).strip(),
            "Q_CHOICE1": safe_text(item.get("a", "")).strip(),
            "Q_CHOICE2": safe_text(item.get("b", "")).strip(),
            "Q_CHOICE3": safe_text(item.get("c", "")).strip(),
            "Q_CHOICE4": safe_text(item.get("d", "")).strip(),
        }
        apply_scalar_map(cells[idx - 1], mapping)

    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.add_run("MODEL ANSWER KEY")
    key_table = doc.add_table(rows=25, cols=2)
    key_table.style = question_table.style
    for row_index in range(25):
        left = row_index + 1
        right = row_index + 26
        key_table.rows[row_index].cells[0].text = f"{left}. {safe_text(mcqs[left - 1].get('answer', '')).strip().upper()}"
        key_table.rows[row_index].cells[1].text = f"{right}. {safe_text(mcqs[right - 1].get('answer', '')).strip().upper()}"

    normalize_all_runs_font(doc)
    return doc


def render_tos_snapshot_page(image_path, *, term: str, course_code: str, course_title: str) -> Document:
    doc = Document()
    enforce_document_font(doc)
    heading = doc.add_paragraph()
    run = heading.add_run(f"{term} TOS SNAPSHOT")
    run.bold = True
    set_run_font(run)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"{course_code} - {course_title}")
    set_run_font(subtitle_run)

    doc.add_picture(str(image_path), width=Inches(7.2))
    normalize_all_runs_font(doc)
    return doc


_ONES = {
    0: "zero",
    1: "one",
    2: "two",
    3: "three",
    4: "four",
    5: "five",
    6: "six",
    7: "seven",
    8: "eight",
    9: "nine",
    10: "ten",
    11: "eleven",
    12: "twelve",
    13: "thirteen",
    14: "fourteen",
    15: "fifteen",
    16: "sixteen",
    17: "seventeen",
    18: "eighteen",
    19: "nineteen",
}

_TENS = {
    20: "twenty",
    30: "thirty",
    40: "forty",
    50: "fifty",
    60: "sixty",
    70: "seventy",
    80: "eighty",
    90: "ninety",
}


def int_to_words(value: int) -> str:
    if value < 0:
        return f"minus {int_to_words(-value)}"
    if value < 20:
        return _ONES[value]
    if value < 100:
        tens = (value // 10) * 10
        remainder = value % 10
        return _TENS[tens] if remainder == 0 else f"{_TENS[tens]}-{_ONES[remainder]}"
    if value < 1000:
        hundreds = value // 100
        remainder = value % 100
        head = f"{_ONES[hundreds]} hundred"
        return head if remainder == 0 else f"{head} {int_to_words(remainder)}"
    if value < 1_000_000:
        thousands = value // 1000
        remainder = value % 1000
        head = f"{int_to_words(thousands)} thousand"
        return head if remainder == 0 else f"{head} {int_to_words(remainder)}"
    raise ValueError(f"Page count too large to format: {value}")


def format_page_count(value: int) -> str:
    return f"{int_to_words(value)} ({value})"


def compute_docx_page_count(docx_path: Path) -> int:
    command = f"""
$ErrorActionPreference = 'Stop'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open('{str(docx_path.resolve()).replace("'", "''")}', $false, $true)
  $pages = $doc.ComputeStatistics(2)
  Write-Output $pages
}} finally {{
  if ($doc -ne $null) {{
    try {{ $doc.Close($false) }} catch {{}}
  }}
  if ($word -ne $null) {{
    try {{ $word.Quit() }} catch {{}}
  }}
}}
""".strip()
    result = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-Command", command],
        capture_output=True,
        text=True,
        timeout=300,
    )
    if result.returncode != 0:
        if result.stdout:
            print(result.stdout, file=sys.stderr)
        if result.stderr:
            print(result.stderr, file=sys.stderr)
        raise RuntimeError(f"Failed to compute page count for {docx_path}")
    output = (result.stdout or "").strip().splitlines()
    if not output:
        raise RuntimeError(f"Word did not return a page count for {docx_path}")
    try:
        pages = int(output[-1].strip())
    except ValueError as exc:
        raise RuntimeError(f"Unexpected page count output for {docx_path}: {result.stdout!r}") from exc
    return max(pages, 1)


def build_ia_term_doc(output_path: Path, parts: list[Document], *, temp_dir: Path) -> None:
    base_path = temp_dir / "000_ia_base.docx"
    parts[0].save(base_path)
    base_doc = Document(base_path)
    composer = Composer(base_doc)
    for index, part in enumerate(parts[1:], start=1):
        breaker = composer.doc.add_paragraph().add_run()
        breaker.add_break(WD_BREAK.PAGE)
        part_path = temp_dir / f"part_{index:03d}.docx"
        part.save(part_path)
        composer.append(Document(part_path))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))


def assemble_ia_term(payload_path, output_path, templates_dir="templates/IA TEMPLATES"):
    payload_path = Path(payload_path)
    output_path = Path(output_path)
    templates_dir = Path(templates_dir)

    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    term = _get_term(payload)
    exams = payload.get("exams") or {}
    if not isinstance(exams, dict):
        raise ValueError("IA payload missing required field: exams")

    course_code = safe_text(exams.get("course_code", "")).strip()
    course_title = safe_text(exams.get("course_title", "")).strip() or safe_text(payload.get("qualification_title", "")).strip()
    mcq_key = "midterm_mcqs" if term == "MIDTERM" else "finals_mcqs"
    template_key = "midterm" if term == "MIDTERM" else "finals"
    exam_term = "MIDTERM" if term == "MIDTERM" else "FINAL"
    mcqs = exams.get(mcq_key)
    tos_snapshot_path = safe_text(payload.get("tos_snapshot_path", "")).strip()

    if not course_code:
        raise ValueError("IA payload exams.course_code must not be empty")
    if not isinstance(mcqs, list):
        raise ValueError(f"IA payload exams.{mcq_key} must be present (50 MCQs)")

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        final_page_count = ""
        for _ in range(3):
            working_payload = dict(payload)
            if final_page_count:
                working_payload["page_count"] = final_page_count

            parts = [
                render_front_page(templates_dir / IA_TEMPLATE_FILES["front_page"], working_payload),
                render_ia_plan(templates_dir / IA_TEMPLATE_FILES["ia_plan"], working_payload),
                render_specific_instructions(templates_dir / IA_TEMPLATE_FILES["specific_instructions"], working_payload),
                render_oral_questions(templates_dir / IA_TEMPLATE_FILES["oral_questions"], working_payload),
            ]
            if tos_snapshot_path:
                parts.append(
                    render_tos_snapshot_page(
                        tos_snapshot_path,
                        term=term,
                        course_code=course_code,
                        course_title=course_title,
                    )
                )
            parts.append(
                render_exam(
                    templates_dir / IA_TEMPLATE_FILES[template_key],
                    course_code=course_code,
                    course_title=course_title,
                    term=exam_term,
                    mcqs=mcqs,
                )
            )

            build_ia_term_doc(output_path, parts, temp_dir=temp_dir)
            measured = format_page_count(compute_docx_page_count(output_path))
            if measured == final_page_count:
                break
            final_page_count = measured


def main(argv):
    if len(argv) < 3:
        print("Usage: python tools/assemble_ia_term.py <payload.json> <output.docx>", file=sys.stderr)
        return 2
    assemble_ia_term(argv[1], argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
