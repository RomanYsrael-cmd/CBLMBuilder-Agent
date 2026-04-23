import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAME = "Bookman Old Style"
FONT_SIZE = 12


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(v) for v in value)
    return str(value)


def set_run_font(run):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)


def enforce_document_font(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = FONT_NAME
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                style.font.size = Pt(FONT_SIZE)
            except Exception:
                pass


def normalize_placeholder_forms(key):
    return [
        "{{" + key + "}}",
        "{{" + key + ":}}",  # handles accidental template form like {{uc_no:}}
    ]


def split_bold_segments(text):
    """
    Splits text into segments:
    [
      ("plain text", False),
      ("bold text", True),
      ...
    ]
    Supports <strong>...</strong> and <b>...</b>
    """
    pattern = re.compile(r"(<strong>.*?</strong>|<b>.*?</b>)", re.IGNORECASE | re.DOTALL)
    parts = re.split(pattern, text)
    segments = []

    for part in parts:
        if not part:
            continue

        strong_match = re.fullmatch(r"<strong>(.*?)</strong>", part, re.IGNORECASE | re.DOTALL)
        b_match = re.fullmatch(r"<b>(.*?)</b>", part, re.IGNORECASE | re.DOTALL)

        if strong_match:
            segments.append((strong_match.group(1), True))
        elif b_match:
            segments.append((b_match.group(1), True))
        else:
            segments.append((part, False))

    return segments


def clear_paragraph(paragraph):
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)


def write_formatted_text_to_paragraph(paragraph, text):
    """
    Writes text into a paragraph while supporting:
    - line breaks
    - <strong>...</strong>
    - <b>...</b>
    """
    clear_paragraph(paragraph)
    first_segment = True

    for segment_text, is_bold in split_bold_segments(text):
        lines = segment_text.split("\n")

        for i, line in enumerate(lines):
            if not first_segment and i == 0:
                pass

            run = paragraph.add_run(line)
            run.bold = is_bold
            set_run_font(run)

            if i < len(lines) - 1:
                run.add_break()

            first_segment = False


def replace_placeholder_in_paragraph(paragraph, placeholder, replacement_text):
    if placeholder not in paragraph.text:
        return

    full_text = paragraph.text.replace(placeholder, replacement_text)
    write_formatted_text_to_paragraph(paragraph, full_text)


def replace_in_paragraphs(paragraphs, placeholder, value):
    replacement_text = safe_text(value)
    for paragraph in paragraphs:
        replace_placeholder_in_paragraph(paragraph, placeholder, replacement_text)


def replace_in_table(table, placeholder, value):
    for row in table.rows:
        for cell in row.cells:
            replace_everywhere(cell, placeholder, value)


def replace_everywhere(doc_part, placeholder, value):
    replace_in_paragraphs(doc_part.paragraphs, placeholder, value)
    for table in doc_part.tables:
        replace_in_table(table, placeholder, value)


def apply_payload(doc, payload):
    for key, value in payload.items():
        for placeholder in normalize_placeholder_forms(key):
            replace_everywhere(doc, placeholder, value)


def normalize_all_runs_font(doc_part):
    for paragraph in doc_part.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)

    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_all_runs_font(cell)


def fill_docx(template_path, output_path, payload):
    doc = Document(template_path)

    enforce_document_font(doc)
    apply_payload(doc, payload)
    normalize_all_runs_font(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    if len(sys.argv) != 4:
        print("Usage: run_fill_cblm.ps1 <template.docx> <payload.json> <output.docx>")
        sys.exit(1)

    template_path = Path(sys.argv[1])
    payload_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    with payload_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)

    fill_docx(template_path, output_path, payload)
    print(str(output_path))


if __name__ == "__main__":
    main()
