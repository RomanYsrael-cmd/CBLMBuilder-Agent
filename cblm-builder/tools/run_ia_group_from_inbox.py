import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docxcompose.composer import Composer


INBOX_DIR = Path("inbox")
PROCESSED_DIR = Path("processed") / "cblm"
STATE_DIR = Path("state") / "ia_payloads"
OUTPUT_DIR = Path("output") / "ia"


LO_HEAD_RE = re.compile(r"^LEARNING OUTCOME NO\.\s*(\d+)\s*-\s*(.+)$", re.I)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _safe_filename(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "_", (text or "").strip())
    text = re.sub(r"_+", "_", text).strip("_")
    return text or "Unit"


def _list_inbox_unit_docx() -> list[Path]:
    if not INBOX_DIR.exists():
        return []
    files = [p for p in INBOX_DIR.glob("CBLM_Unit_*.docx") if p.is_file()]
    return sorted(files, key=lambda p: p.name.lower())


def _parse_filename(path: Path) -> tuple[int, str, str]:
    """
    Expected: CBLM_Unit_<n>_<COURSE_CODE>_<Title>.docx
    COURSE_CODE may contain underscores (e.g., ICC_5, IIPC_13).
    Infer COURSE_CODE from the first 1–2 tokens after the unit number:
    - Prefer <LETTERS>_<DIGITS> (e.g., ICC_5, IIPC_13, IS_6, IPE_10)
    - Otherwise accept <LETTERS><DIGITS> (e.g., ICC14)
    """
    m = re.match(r"^CBLM_Unit_(\d+)_(.+)\.docx$", path.name, re.I)
    if not m:
        raise ValueError(f"Invalid filename (expected CBLM_Unit_<n>_*.docx): {path.name}")
    unit_index = int(m.group(1))
    remainder = m.group(2)
    tokens = remainder.split("_")
    if len(tokens) < 3:
        raise ValueError(f"Invalid filename (expected CBLM_Unit_<n>_<COURSE_CODE>_<Title>.docx): {path.name}")
    course_code = ""
    title_tokens: list[str] = []

    if len(tokens) >= 2 and re.match(r"^[A-Za-z]+_[0-9]+$", f"{tokens[0]}_{tokens[1]}"):
        course_code = f"{tokens[0]}_{tokens[1]}"
        title_tokens = tokens[2:]
    elif re.match(r"^[A-Za-z]+[0-9]+$", tokens[0]):
        course_code = tokens[0]
        title_tokens = tokens[1:]
    else:
        course_code = tokens[0]
        title_tokens = tokens[1:]

    title = "_".join(title_tokens)
    return unit_index, course_code, title


def _extract_unit_of_competency(paragraphs: list[str]) -> str:
    for t in paragraphs:
        if "The unit of competency" not in t:
            continue
        m = re.search(r"The unit of competency,\s*(.+?),\s*is\s+one\s+of", t)
        if not m:
            continue
        cand = _norm(m.group(1))
        return cand.strip("\"'“”‘’").strip()
    return ""


def _extract_module_title(paragraphs: list[str]) -> str:
    for t in paragraphs:
        m = re.search(r"The module,\s*(.+?),\s*contains", t)
        if m:
            return _norm(m.group(1))
    return ""


def _extract_next_unit(paragraphs: list[str]) -> str:
    for t in paragraphs:
        m = re.search(
            r"You need to complete this module before you can perform the module on,\s*(.+)$",
            t,
        )
        if m:
            return _norm(m.group(1))
    return ""


def _extract_front_matter_from_tables(doc: Document) -> dict:
    """
    Best-effort extraction from the front-matter tables that contain labels like:
    Sector / Qualification Title / Unit of Competency / Module Title.
    """
    out = {"qualification_title": "", "unit_of_competency": "", "module_title": ""}

    def cell_text(cell) -> str:
        return _norm(getattr(cell, "text", "") or "")

    labelish = {
        "sector",
        "qualification title",
        "qualification",
        "unit of competency",
        "module title",
        "module descriptor",
        "prepared by",
        "code",
        "no.",
    }

    def is_label(value: str) -> bool:
        low = _norm(value).lower()
        if not low:
            return False
        if low in ("code", "module title", "unit of competency", "qualification title", "sector"):
            return True
        return any(k in low for k in labelish)

    def pick_value(texts: list[str], start_idx: int) -> str:
        for cand in texts[start_idx + 1 :]:
            cand = _norm(cand)
            if not cand:
                continue
            if is_label(cand):
                continue
            return cand
        for cand in reversed(texts):
            cand = _norm(cand)
            if cand and not is_label(cand):
                return cand
        return ""

    for table in doc.tables:
        for row in table.rows:
            texts = [cell_text(c) for c in row.cells]
            for idx, t in enumerate(texts):
                low = t.lower()
                if "qualification title" in low and idx + 1 < len(texts):
                    out["qualification_title"] = out["qualification_title"] or pick_value(texts, idx)
                if "unit of competency" in low and idx + 1 < len(texts):
                    out["unit_of_competency"] = out["unit_of_competency"] or pick_value(texts, idx)
                if "module title" in low and idx + 1 < len(texts):
                    out["module_title"] = out["module_title"] or pick_value(texts, idx)

    return out


def _extract_qualification_title(paragraphs: list[str], module_title: str) -> str:
    """
    Heuristic:
    - Prefer early standalone-ish title lines (not labels, not sentences).
    - Fallback to module_title without parentheticals.
    """
    skip_prefixes = (
        "cb lm",
        "cblm",
        "competency-based learning materials",
        "competency based learning materials",
        "how to use this competency",
        "how to use this competency- based",
        "how to use this competency-",
        "module content",
        "trainee",
        "institution",
        "trainer",
        "qualification",
        "unit of competency",
        "module title",
    )
    for t in paragraphs[:60]:
        if not t:
            continue
        lower = t.lower()
        if any(lower.startswith(p) for p in skip_prefixes):
            continue
        if lower in ("competency-based learning materials", "competency based learning materials"):
            continue
        if len(t.split()) <= 8 and not t.endswith(".") and ":" not in t:
            return t
    if module_title:
        return module_title.split("(")[0].strip()
    return ""


def _extract_learning_outcomes(doc: Document) -> list[dict]:
    los: list[dict] = []
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        t = _norm(paragraphs[i].text)
        m = LO_HEAD_RE.match(t)
        if not m:
            i += 1
            continue

        lo = {"index": len(los) + 1, "title": _norm(m.group(2)), "contents": []}

        j = i + 1
        while j < len(paragraphs) and not _norm(paragraphs[j].text).lower().startswith("contents"):
            j += 1

        if j < len(paragraphs) and _norm(paragraphs[j].text).lower().startswith("contents"):
            k = j + 1
            while k < len(paragraphs):
                pk = paragraphs[k]
                tk = _norm(pk.text)
                if not tk:
                    k += 1
                    continue
                if LO_HEAD_RE.match(tk):
                    break
                if tk.lower().startswith("assessment criteria"):
                    break

                style = (getattr(pk.style, "name", "") or "").lower()
                if style.startswith("list"):
                    lo["contents"].append({"index": len(lo["contents"]) + 1, "title": tk})
                    k += 1
                    continue
                break

        los.append(lo)
        i = j + 1

    return los


def _tokens(text: str) -> list[str]:
    return [t for t in re.split(r"[^A-Za-z0-9]+", (text or "").lower()) if t]


def _norm_q(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()


def _require_keyword_coverage(text: str, anchors: list[str], min_hits: int = 1) -> bool:
    hay = " " + (text or "").lower() + " "
    hits = 0
    for a in anchors:
        a = (a or "").strip()
        if not a:
            continue
        # use a few strong tokens from the anchor
        toks = [t for t in _tokens(a) if len(t) >= 4]
        if not toks:
            continue
        if any((" " + t + " ") in hay for t in toks[:5]):
            hits += 1
    return hits >= min_hits


def _build_ia_block(payload: dict, *, unit_index: int, prior_question_texts: list[str], used_project_anchors: set[str]) -> dict:
    """
    Build IA-only fields from the extracted unit outline so each unit's project,
    oral questions, and acceptable answers are derived from that unit's content.
    """
    unit = payload["current_unit"]
    qualification_title = payload.get("qualification_title", "")
    module_title = unit.get("module_title", "") or qualification_title
    unit_of_competency = (unit.get("unit_of_competency", "") or "").strip()

    def iter_contents():
        for lo in unit.get("learning_outcomes", []) or []:
            for content in lo.get("contents", []) or []:
                title = _norm(content.get("title", ""))
                if title:
                    yield lo, content, title

    contents = list(iter_contents())
    lo_titles = [
        _norm(lo.get("title", ""))
        for lo in unit.get("learning_outcomes", []) or []
        if _norm(lo.get("title", ""))
    ]

    def pick_titles(count: int, offset: int = 0) -> list[str]:
        seen: set[str] = set()
        picked: list[str] = []
        ordered = [t for _, _, t in contents]
        if ordered and offset:
            offset = offset % len(ordered)
            ordered = ordered[offset:] + ordered[:offset]
        for title in ordered:
            key = title.lower()
            if key in seen:
                continue
            seen.add(key)
            picked.append(title)
            if len(picked) >= count:
                break
        while len(picked) < count and lo_titles:
            picked.append(lo_titles[len(picked) % len(lo_titles)])
        while len(picked) < count:
            picked.append(unit_of_competency or module_title or qualification_title or "this unit")
        return picked

    # Guardrail strategy:
    # - anchor project and questions to content titles (rotate per unit)
    # - ensure project anchor differs across units in the same group
    # - avoid exact-duplicate questions across units (normalized)
    base_offset = unit_index * 3
    picks = pick_titles(5, offset=base_offset)
    project_core = _norm(unit_of_competency or module_title or qualification_title or "Institutional Assessment")
    project_name = f"{project_core} Evidence Portfolio: {picks[0]}"

    instructions = [
        "1. Review the IA plan and identify the required evidence for each Learning Outcome and content item.",
        f"2. Create a one-page overview for the unit of competency: {project_core}.",
        f"3. Prepare evidence artifacts for at least three contents (e.g., short notes, diagrams, examples, or outputs), prioritizing: {picks[0]}, {picks[1]}, and {picks[2]}.",
        "4. Organize your evidence by LO number and content index, and label each artifact clearly.",
        "5. Perform the demonstration as directed by the assessor and explain the purpose of each artifact.",
        "6. Answer the oral questions by referencing your evidence and using correct technical terms from the unit.",
        "7. Participate in the interview and justify the steps/decisions you made, including any checks you performed to ensure correctness.",
        "8. Submit your final evidence package in the required format (printed or digital) for recording and evaluation.",
    ]

    materials = [
        "Pen and paper / notebook",
        "Laptop/desktop computer",
        "Printed or digital copy of the CBLM unit and IA plan",
        "Basic office supplies (highlighters, sticky notes) for organizing evidence",
    ]
    if any("diagram" in t.lower() or "model" in t.lower() or "flow" in t.lower() for t in picks):
        materials.append("Diagramming tool (e.g., draw.io, Visio, or equivalent)")

    def make_oral_questions(p: list[str]) -> list[dict]:
        return [
            {
                "question": f"Explain {p[0]} and give one practical workplace example of its use.",
                "acceptable_answer": f"Correctly describes {p[0]} and provides a realistic workplace example aligned to the unit context.",
            },
            {
                "question": f"Describe the main steps or process you would follow to address {p[1]}.",
                "acceptable_answer": f"Outlines a logical sequence of steps for {p[1]} and mentions checks/validation to ensure correctness.",
            },
            {
                "question": f"What is one common mistake related to {p[2]} and how would you prevent it?",
                "acceptable_answer": f"Identifies a plausible error for {p[2]} and gives a practical prevention method (checklist, review, validation, documentation).",
            },
            {
                "question": f"Differentiate {p[3]} from {p[4]} based on purpose or application.",
                "acceptable_answer": f"States a clear difference between {p[3]} and {p[4]} and explains when each would be used.",
            },
            {
                "question": "How do you ensure your evidence/output meets the required performance criteria for this unit?",
                "acceptable_answer": "Explains checking against the IA plan/criteria, validating outputs for completeness and correctness, and documenting evidence clearly.",
            },
        ]

    oral_questions = make_oral_questions(picks)
    # Guardrails: enforce uniqueness and anchor coverage; retry a few times with different offsets if needed.
    def questions_ok(qs: list[dict]) -> bool:
        texts = [q.get("question", "") for q in qs]
        # Allow the final generic QA question to repeat across units; enforce uniqueness on the first 4 only.
        texts_strict = texts[:4]
        # within-unit duplicates / near-duplicates
        for i in range(len(texts_strict)):
            for j in range(i + 1, len(texts_strict)):
                if _norm_q(texts_strict[i]) == _norm_q(texts_strict[j]):
                    return False
        # cross-unit duplicates / near-duplicates
        for t in texts_strict:
            for prev in prior_question_texts:
                if _norm_q(t) == _norm_q(prev):
                    return False
        # keyword coverage: at least 3 questions mention a strong token from at least one content anchor
        anchors = picks[:4]
        covered = sum(1 for t in texts_strict if _require_keyword_coverage(t, anchors, min_hits=1))
        if covered < 3:
            return False
        return True

    attempt = 0
    while attempt < 3 and (picks[0].lower() in used_project_anchors or not questions_ok(oral_questions)):
        attempt += 1
        picks = pick_titles(5, offset=base_offset + (attempt * 7))
        oral_questions = make_oral_questions(picks)
        project_name = f"{project_core} Evidence Portfolio: {picks[0]}"
    if picks[0].lower() in used_project_anchors or not questions_ok(oral_questions):
        raise RuntimeError("IA guardrails failed: could not produce unique, content-anchored oral questions/project for this unit.")
    used_project_anchors.add(picks[0].lower())

    # Ensure only question/acceptable_answer keys are present (template expects this shape).
    oral_questions = [{"question": q["question"], "acceptable_answer": q["acceptable_answer"]} for q in oral_questions]

    interview_questions = [
        f"Walk through your evidence portfolio for {project_core} and explain how it demonstrates competence.",
        f"Which content did you find most challenging ({picks[0]} / {picks[1]} / {picks[2]}) and how did you address it?",
        "What checks did you perform to ensure accuracy and consistency of your outputs/evidence?",
        "If you had more time, how would you improve the quality or clarity of your evidence package?",
        "Describe one professional or ethical consideration relevant to this unit and how you applied it during the activity.",
    ]

    return {
        "name_of_project": project_name,
        "Specific_Instructions": "\n".join(instructions),
        "instructions_for_demo": (
            f"Demonstrate the required tasks for {project_core} using your prepared evidence artifacts. "
            "Explain each step clearly and answer follow-up questions from the assessor."
        ),
        "list_of_materials_and_equipment": "\n".join(materials),
        "oral_questions": oral_questions,
        "interview_questions": interview_questions,
    }


def _assemble_ia(payload_path: Path, output_path: Path) -> None:
    python_exe = Path(".venv") / "Scripts" / "python.exe"
    cmd = [str(python_exe), str(Path("tools") / "assemble_ia.py"), str(payload_path), str(output_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError((result.stdout or "") + (result.stderr or ""))


def _merge_docx(inputs: list[Path], output_path: Path) -> None:
    python_exe = Path(".venv") / "Scripts" / "python.exe"
    cmd = [str(python_exe), str(Path("tools") / "merge_docx.py"), *[str(p) for p in inputs], str(output_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError((result.stdout or "") + (result.stderr or ""))


def main(argv: list[str]) -> int:
    unit_files = _list_inbox_unit_docx()
    if not unit_files:
        return 0

    parsed: list[tuple[int, str, str, Path]] = []
    for p in unit_files:
        try:
            unit_index, course_code, title = _parse_filename(p)
        except ValueError as e:
            failed_dir = Path("failed")
            failed_dir.mkdir(parents=True, exist_ok=True)
            dest = failed_dir / p.name
            if dest.exists():
                dest = failed_dir / f"{p.stem}__INVALID{p.suffix}"
            try:
                shutil.move(str(p), str(dest))
            except Exception:
                pass
            sys.stderr.write(str(e) + "\n")
            return 2
        parsed.append((unit_index, course_code, title, p))

    course_codes = sorted({course_code for _, course_code, _, _ in parsed})
    chosen_course = course_codes[0]

    chosen_units = [(unit_index, title, p) for unit_index, course_code, title, p in parsed if course_code == chosen_course]
    chosen_units.sort(key=lambda x: x[0])
    if not chosen_units:
        sys.stderr.write(f"No unit files found for chosen course_code: {chosen_course}\n")
        return 2

    out_dir = OUTPUT_DIR / chosen_course
    state_dir = STATE_DIR / chosen_course
    processed_dir = PROCESSED_DIR / chosen_course
    out_dir.mkdir(parents=True, exist_ok=True)
    state_dir.mkdir(parents=True, exist_ok=True)
    processed_dir.mkdir(parents=True, exist_ok=True)

    unit_outputs: list[Path] = []
    to_move_inputs: list[Path] = []
    payload_paths: list[Path] = []
    prior_oral_questions: list[str] = []
    used_project_anchors: set[str] = set()

    for unit_index, title, input_path in chosen_units:
        doc = Document(str(input_path))
        head_paras = [_norm(p.text) for p in doc.paragraphs[:250] if _norm(p.text)]

        fm = _extract_front_matter_from_tables(doc)
        module_title = fm.get("module_title") or ""
        unit_of_competency = fm.get("unit_of_competency") or ""
        qual_from_tables = fm.get("qualification_title") or ""

        # Sanity-check table-derived values; fall back to paragraph extraction when table parsing misaligns.
        if not module_title or module_title.lower() in ("code", "module title"):
            module_title = _extract_module_title(head_paras)
        if not unit_of_competency or unit_of_competency.lower() in ("module title", "unit of competency", "code"):
            unit_of_competency = _extract_unit_of_competency(head_paras)

        # Qualification title: prefer table value; fallback to filename title (more reliable than paragraph heuristics).
        qual = qual_from_tables or title.replace("_", " ").strip()
        # If the extracted value is generic, fall back to the filename title (most reliable in this repo).
        if qual.lower() in (
            "competency-based learning materials",
            "competency based learning materials",
            "how to use this competency- based learning materials",
            "how to use this competency-based learning materials",
        ):
            qual = title.replace("_", " ").strip()
        payload = {
            "qualification_title": qual,
            "current_unit": {
                "unit_of_competency": unit_of_competency,
                "module_title": module_title,
                "next_unit_of_competency": _extract_next_unit(head_paras),
                "learning_outcomes": _extract_learning_outcomes(doc),
                "ia": {},
            },
        }
        payload["current_unit"]["ia"] = _build_ia_block(
            payload,
            unit_index=unit_index,
            prior_question_texts=prior_oral_questions,
            used_project_anchors=used_project_anchors,
        )
        prior_oral_questions.extend([q.get("question", "") for q in (payload["current_unit"]["ia"].get("oral_questions") or [])[:4]])

        safe_short_title = _safe_filename(title)
        payload_path = state_dir / f"IA_Unit_{unit_index}_{safe_short_title}.json"
        output_path = out_dir / f"IA_Unit_{unit_index}_{safe_short_title}.docx"

        payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        _assemble_ia(payload_path, output_path)

        payload_paths.append(payload_path)
        unit_outputs.append(output_path)
        to_move_inputs.append(input_path)

    full_path = out_dir / "IA_FULL.docx"
    tmp_full = out_dir / "IA_FULL__tmp.docx"
    if len(unit_outputs) == 1:
        shutil.copyfile(str(unit_outputs[0]), str(tmp_full))
    else:
        _merge_docx(unit_outputs, tmp_full)
    try:
        if full_path.exists():
            full_path.unlink()
        shutil.move(str(tmp_full), str(full_path))
    except PermissionError as e:
        raise RuntimeError(f"Cannot overwrite {full_path} (file may be open). Close it and rerun. Details: {e}")
    finally:
        try:
            if tmp_full.exists():
                tmp_full.unlink()
        except Exception:
            pass

    # If course-level exams exist (generated after all UCs), append them to IA_FULL.
    _maybe_append_course_exams(chosen_course, out_dir=out_dir)

    processed_inputs: list[Path] = []
    for input_path in to_move_inputs:
        dest = processed_dir / input_path.name
        if dest.exists():
            dest = processed_dir / f"{input_path.stem}__DUPLICATE{input_path.suffix}"
        shutil.move(str(input_path), str(dest))
        processed_inputs.append(dest)

    summary = {
        "course_code": chosen_course,
        "units_processed": [p.name for p in processed_inputs],
        "payloads_written": [p.name for p in payload_paths],
        "ia_unit_outputs": [p.name for p in unit_outputs],
        "ia_full_path": str(full_path),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


def _maybe_append_course_exams(course_code: str, *, out_dir: Path) -> None:
    """
    If course-level exams exist (generated after all UCs), append them to IA_FULL.docx.
    """
    ia_full = out_dir / "IA_FULL.docx"
    if not ia_full.exists():
        return

    mid = out_dir / "03_midterm.docx"
    fin = out_dir / "04_finals.docx"
    if not mid.exists() or not fin.exists():
        return

    tmp = out_dir / "IA_FULL__with_exams__tmp.docx"
    base_doc = Document(str(ia_full))
    composer = Composer(base_doc)
    for p in [mid, fin]:
        composer.doc.add_page_break()
        composer.append(Document(str(p)))
    try:
        if tmp.exists():
            tmp.unlink()
        composer.save(str(tmp))
        ia_full.unlink()
        tmp.rename(ia_full)
    except PermissionError as e:
        raise RuntimeError(f"Cannot overwrite {ia_full} (file may be open). Close it and rerun. Details: {e}")


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
