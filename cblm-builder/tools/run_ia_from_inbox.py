import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from ia_oral_questions import build_oral_questions_from_payload


INBOX_DIR = Path("inbox")
PROCESSED_DIR = Path("processed") / "cblm"
STATE_DIR = Path("state") / "ia_payloads"
OUTPUT_DIR = Path("output") / "ia"


LO_HEAD_RE = re.compile(r"^LEARNING OUTCOME NO\.\s*(\d+)\s*-\s*(.+)$", re.I)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _safe_filename(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "_", (text or "").strip())
    text = re.sub(r"_+", "_", text).strip("_")
    return text or "Unit"


def _pick_unit_files() -> list[Path]:
    if not INBOX_DIR.exists():
        return []
    files = [p for p in INBOX_DIR.glob("CBLM_Unit_*.docx") if p.is_file()]
    return sorted(files, key=lambda p: p.name.lower())


def _extract_unit_index_and_short_title(path: Path) -> tuple[int, str]:
    m = re.match(r"^CBLM_Unit_(\d+)_(.+)\.docx$", path.name, re.I)
    if not m:
        return 1, _safe_filename(path.stem)
    return int(m.group(1)), _safe_filename(m.group(2))


def _extract_unit_of_competency(paragraphs: list[str]) -> str:
    for t in paragraphs:
        if "The unit of competency" not in t:
            continue
        m = re.search(r"The unit of competency,\s*(.+?),\s*is\s+one\s+of", t)
        if not m:
            continue
        cand = _norm(m.group(1))
        # strip straight + smart quotes if present
        return cand.strip("\"'“”‘’").strip()
    return ""


def _extract_module_title(paragraphs: list[str]) -> str:
    for t in paragraphs:
        m = re.search(r"The module,\s*(.+?),\s*contains", t)
        if m:
            return _norm(m.group(1))
    return ""


def _extract_next_unit(paragraphs: list[str]) -> str:
    for t in paragraphs:
        m = re.search(
            r"You need to complete this module before you can perform the module on,\s*(.+)$",
            t,
        )
        if m:
            return _norm(m.group(1))
    return ""


def _extract_qualification_title(paragraphs: list[str], module_title: str) -> str:
    # Prefer a standalone title line. Fallback to module prefix.
    for t in paragraphs:
        if not t:
            continue
        if t == "Information Management":
            return t
    if module_title:
        return module_title.split("(")[0].strip()
    return ""


def _extract_learning_outcomes(doc: Document) -> list[dict]:
    los: list[dict] = []
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        t = _norm(paragraphs[i].text)
        m = LO_HEAD_RE.match(t)
        if not m:
            i += 1
            continue

        lo = {"index": len(los) + 1, "title": _norm(m.group(2)), "contents": []}

        j = i + 1
        while j < len(paragraphs) and not _norm(paragraphs[j].text).lower().startswith("contents"):
            j += 1

        if j < len(paragraphs) and _norm(paragraphs[j].text).lower().startswith("contents"):
            k = j + 1
            while k < len(paragraphs):
                pk = paragraphs[k]
                tk = _norm(pk.text)
                if not tk:
                    k += 1
                    continue
                if LO_HEAD_RE.match(tk):
                    break
                if tk.lower().startswith("assessment criteria"):
                    break

                style = (getattr(pk.style, "name", "") or "").lower()
                if style.startswith("list"):
                    lo["contents"].append({"index": len(lo["contents"]) + 1, "title": tk})
                    k += 1
                    continue
                break

        los.append(lo)
        i = j + 1

    return los


def _build_ia_block(payload: dict) -> dict:
    unit = payload["current_unit"]
    qualification_title = payload.get("qualification_title", "")
    module_title = unit.get("module_title", "") or qualification_title

    return {
        "name_of_project": f"Institutional Assessment Activity ({module_title})",
        "Specific_Instructions": "\n".join(
            [
                "1. Read the assessment brief and confirm the allotted time with the trainer/assessor.",
                "2. Review the learning outcomes and contents listed in the IA plan for this unit.",
                "3. Prepare the required evidence for each content item (notes, diagrams, examples, or outputs) based on your trainer’s instructions.",
                "4. Perform the demonstration/work activity as directed and observe workplace-safe practices.",
                "5. Present your outputs/evidence to the assessor and answer the oral questions.",
                "6. Participate in the interview and clarify any steps or decisions you made during the activity.",
                "7. Submit final outputs in the required format (printed or digital) for recording and evaluation.",
            ]
        ),
        "instructions_for_demo": (
            "Demonstrate the required tasks for the unit based on the IA plan, explain your steps clearly, "
            "and answer follow-up questions from the assessor."
        ),
        "list_of_materials_and_equipment": "\n".join(
            [
                "Pen and paper / notebook",
                "Laptop/desktop computer (if required by the activity)",
                "Access to references or handouts provided by the trainer",
            ]
        ),
        "oral_questions": build_oral_questions_from_payload(payload),
        "interview_questions": [
            "Explain the reasoning behind the steps you followed during the activity.",
            "What assumptions did you make, and how would you verify them in an actual workplace?",
            "How would you improve your process to reduce time and errors while maintaining quality?",
            "What documentation would you keep as evidence of correct performance?",
            "Describe one ethical or professional consideration relevant to this unit.",
        ],
    }


def main(argv: list[str]) -> int:
    unit_files = _pick_unit_files()
    if not unit_files:
        return 0

    if len(unit_files) > 1 and (len(argv) < 2 or argv[1].strip().lower() != "continue"):
        chosen = unit_files[0].name
        print(f"Multiple CBLM Unit DOCX files found. Next file (deterministic): {chosen}")
        print("Reply with `continue` to process it.")
        return 0

    chosen_path = unit_files[0]
    unit_index, short_title = _extract_unit_index_and_short_title(chosen_path)

    doc = Document(str(chosen_path))
    head_paras = [_norm(p.text) for p in doc.paragraphs[:200] if _norm(p.text)]

    module_title = _extract_module_title(head_paras)
    payload = {
        "qualification_title": _extract_qualification_title(head_paras, module_title),
        "current_unit": {
            "unit_of_competency": _extract_unit_of_competency(head_paras),
            "module_title": module_title,
            "next_unit_of_competency": _extract_next_unit(head_paras),
            "learning_outcomes": _extract_learning_outcomes(doc),
            "ia": {},
        },
    }
    payload["current_unit"]["ia"] = _build_ia_block(payload)

    STATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)

    payload_path = STATE_DIR / f"IA_Unit_{unit_index}_{short_title}.json"
    output_path = OUTPUT_DIR / f"IA_Unit_{unit_index}_{short_title}.docx"

    payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    # Assemble IA docx (same environment/venv).
    import subprocess

    result = subprocess.run(
        [str(Path(".venv") / "Scripts" / "python.exe"), str(Path("tools") / "assemble_ia.py"), str(payload_path), str(output_path)],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        sys.stderr.write(result.stdout)
        sys.stderr.write(result.stderr)
        return result.returncode

    # Move processed input out of inbox.
    dest = PROCESSED_DIR / chosen_path.name
    if dest.exists():
        dest = PROCESSED_DIR / f"{chosen_path.stem}__DUPLICATE{chosen_path.suffix}"
    shutil.move(str(chosen_path), str(dest))

    print(f"input: {chosen_path}")
    print(f"payload: {payload_path}")
    print(f"output: {output_path}")
    print(f"moved_to: {dest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
