import argparse
import re
import sys
from collections import Counter
from itertools import combinations
from pathlib import Path

from docx import Document


def normalize(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def tokenize(text: str) -> list[str]:
    return [t for t in normalize(text).split(" ") if t]


def shingles(tokens: list[str], size: int = 5) -> set[str]:
    if not tokens:
        return set()
    if len(tokens) < size:
        return {" ".join(tokens)}
    return {" ".join(tokens[i : i + size]) for i in range(len(tokens) - size + 1)}


def jaccard(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    return len(left & right) / len(left | right)


def extract_questions(doc_path: Path) -> list[str]:
    doc = Document(doc_path)
    question_table = None
    for table in doc.tables:
        if len(table.rows) == 25 and len(table.columns) == 2:
            question_table = table
            break
    if question_table is None:
        raise ValueError("Could not locate the 25x2 question table in the exam template.")

    questions: list[str] = []
    for row in question_table.rows:
        for cell in row.cells:
            raw = "\n".join(p.text for p in cell.paragraphs).strip()
            if not raw:
                continue
            lines = [line.strip() for line in raw.splitlines() if line.strip()]
            if not lines:
                continue
            question = re.sub(r"^_+\s*", "", lines[0])
            question = re.sub(r"^\s*\d+\s*[\.)]?\s*", "", question)
            questions.append(question.strip())
    return questions


def classify_cognitive_level(question: str) -> str:
    lower = normalize(question)
    if any(
        phrase in lower
        for phrase in [
            "what is",
            "which term",
            "which statement",
            "identify",
            "define",
            "refers to",
            "best describes",
            "is used for",
        ]
    ):
        if "scenario" not in lower and "given" not in lower and "best action" not in lower:
            return "knowledge"

    if any(
        phrase in lower
        for phrase in [
            "why",
            "how does",
            "how do",
            "best explains",
            "most likely explains",
            "interpret",
            "compare",
            "distinguish",
            "relationship between",
        ]
    ):
        return "comprehension"

    if any(
        phrase in lower
        for phrase in [
            "given",
            "a company",
            "a team",
            "a systems analyst",
            "what should",
            "which action",
            "best action",
            "best example",
            "most appropriate",
            "applied",
            "during implementation",
            "working on",
        ]
    ):
        return "application"

    return "unknown"


def stripped_scaffold(question: str) -> str:
    text = question
    text = re.sub(r"\[[^\]]+\]", " ", text)
    text = re.sub(r"\(.*?\)", " ", text)
    text = re.sub(r"'.*?'", " ", text)
    text = re.sub(r'".*?"', " ", text)
    # remove likely topic/title fragments after "under" or "on"
    text = re.sub(r"\bunder\s+[A-Za-z0-9 ,/&\-]+", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\bon\s+[A-Za-z0-9 ,/&\-]+", " ", text, flags=re.IGNORECASE)
    return normalize(text)


def validate_questions(questions: list[str], *, expected_count: int, max_scaffold_count: int) -> None:
    if len(questions) != expected_count:
        raise ValueError(f"Expected {expected_count} questions, found {len(questions)}.")

    tagged = [q for q in questions if re.search(r"\[[^\]]+\]", q)]
    if tagged:
        raise ValueError(f"Found bracketed metadata tags in questions, e.g. '{tagged[0]}'.")

    uc_meta = [q for q in questions if re.search(r"\bUC\s*\d+\b", q, re.IGNORECASE) or "unit of competency" in normalize(q)]
    if uc_meta:
        raise ValueError(f"Found UC/unit metadata in questions, e.g. '{uc_meta[0]}'.")

    normalized = [normalize(q) for q in questions]
    if len(set(normalized)) != len(normalized):
        raise ValueError("Found duplicate question text.")

    prepared = [shingles(tokenize(q)) for q in questions]
    similar_pairs = []
    for i, j in combinations(range(len(questions)), 2):
        score = jaccard(prepared[i], prepared[j])
        if score >= 0.4:
            similar_pairs.append((i + 1, j + 1, score))
    if similar_pairs:
        i, j, score = similar_pairs[0]
        raise ValueError(f"Found near-duplicate questions ({i} vs {j}, similarity={score:.2f}).")

    scaffolds = [stripped_scaffold(q) for q in questions]
    scaffold_counts = Counter(scaffolds)
    common_scaffold, common_count = scaffold_counts.most_common(1)[0]
    if common_scaffold and common_count > max_scaffold_count:
        raise ValueError(f"Too many questions share the same scaffold ({common_count}x): '{common_scaffold[:120]}'")

    stem_patterns = Counter()
    for q in normalized:
        if "refers to which idea" in q:
            stem_patterns["refers_to_idea"] += 1
        if "best example" in q:
            stem_patterns["best_example"] += 1
        if "mainly to accomplish what" in q:
            stem_patterns["mainly_to_accomplish"] += 1
    for pattern, count in stem_patterns.items():
        if count > 6:
            raise ValueError(f"Overused exam stem '{pattern}' appears {count} times.")

    levels = Counter(classify_cognitive_level(q) for q in questions)
    if levels["knowledge"] < 8 or levels["comprehension"] < 8 or levels["application"] < 24:
        raise ValueError(
            "Cognitive-level mix looks off. "
            f"Detected knowledge={levels['knowledge']}, comprehension={levels['comprehension']}, application={levels['application']}."
        )


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate generated exam DOCX quality.")
    parser.add_argument("exam_docx", type=Path)
    parser.add_argument("--expected-count", type=int, default=50)
    parser.add_argument("--max-scaffold-count", type=int, default=4)
    args = parser.parse_args()

    questions = extract_questions(args.exam_docx)
    validate_questions(
        questions,
        expected_count=args.expected_count,
        max_scaffold_count=args.max_scaffold_count,
    )
    print(f"OK: exam quality checks passed for {args.exam_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

