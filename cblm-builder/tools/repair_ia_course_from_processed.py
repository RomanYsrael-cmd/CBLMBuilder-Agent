import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document


PROCESSED_CBLM_DIR = Path("processed") / "cblm"
STATE_DIR = Path("state") / "ia_payloads"
OUTPUT_DIR = Path("output") / "ia"

LO_HEAD_RE = re.compile(r"^LEARNING OUTCOME NO\.\s*(\d+)\s*-\s*(.+)$", re.I)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _safe_filename(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "_", (text or "").strip())
    text = re.sub(r"_+", "_", text).strip("_")
    return text or "Unit"


def _extract_unit_of_competency(paragraphs: list[str]) -> str:
    for t in paragraphs:
        if "The unit of competency" not in t:
            continue
        m = re.search(r"The unit of competency,\s*(.+?),\s*is\s+one\s+of", t)
        if not m:
            continue
        cand = _norm(m.group(1))
        return cand.strip("\"'“”‘’â€œâ€â€˜â€™").strip()
    return ""


def _extract_module_title(paragraphs: list[str]) -> str:
    for t in paragraphs:
        m = re.search(r"The module,\s*(.+?),\s*contains", t)
        if m:
            return _norm(m.group(1))
    return ""


def _extract_next_unit(paragraphs: list[str]) -> str:
    for t in paragraphs:
        m = re.search(
            r"You need to complete this module before you can perform the module on,\s*(.+)$",
            t,
        )
        if m:
            return _norm(m.group(1))
    return ""


def _extract_front_matter_from_tables(doc: Document) -> dict:
    out = {"qualification_title": "", "unit_of_competency": "", "module_title": ""}

    def cell_text(cell) -> str:
        return _norm(getattr(cell, "text", "") or "")

    labelish = {
        "sector",
        "qualification title",
        "qualification",
        "unit of competency",
        "module title",
        "module descriptor",
        "prepared by",
        "code",
        "no.",
    }

    def is_label(value: str) -> bool:
        low = _norm(value).lower()
        if not low:
            return False
        if low in ("code", "module title", "unit of competency", "qualification title", "sector"):
            return True
        return any(k in low for k in labelish)

    def pick_value(texts: list[str], start_idx: int) -> str:
        for cand in texts[start_idx + 1 :]:
            cand = _norm(cand)
            if not cand:
                continue
            if is_label(cand):
                continue
            return cand
        for cand in reversed(texts):
            cand = _norm(cand)
            if cand and not is_label(cand):
                return cand
        return ""

    for table in doc.tables:
        for row in table.rows:
            texts = [cell_text(c) for c in row.cells]
            for idx, t in enumerate(texts):
                low = t.lower()
                if "qualification title" in low and idx + 1 < len(texts):
                    out["qualification_title"] = out["qualification_title"] or pick_value(texts, idx)
                if "unit of competency" in low and idx + 1 < len(texts):
                    out["unit_of_competency"] = out["unit_of_competency"] or pick_value(texts, idx)
                if "module title" in low and idx + 1 < len(texts):
                    out["module_title"] = out["module_title"] or pick_value(texts, idx)
    return out


def _extract_qualification_title(paragraphs: list[str], module_title: str) -> str:
    skip_prefixes = (
        "cb lm",
        "cblm",
        "competency-based learning materials",
        "competency based learning materials",
        "how to use this competency",
        "how to use this competency- based",
        "how to use this competency-",
        "module content",
        "trainee",
        "institution",
        "trainer",
        "qualification",
        "unit of competency",
        "module title",
    )
    for t in paragraphs[:80]:
        if not t:
            continue
        lower = t.lower()
        if any(lower.startswith(p) for p in skip_prefixes):
            continue
        if lower in ("competency-based learning materials", "competency based learning materials"):
            continue
        if len(t.split()) <= 8 and not t.endswith(".") and ":" not in t:
            return t
    if module_title:
        return module_title.split("(")[0].strip()
    return ""


def _extract_learning_outcomes(doc: Document) -> list[dict]:
    los: list[dict] = []
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        t = _norm(paragraphs[i].text)
        m = LO_HEAD_RE.match(t)
        if not m:
            i += 1
            continue

        lo = {"index": len(los) + 1, "title": _norm(m.group(2)), "contents": []}

        j = i + 1
        while j < len(paragraphs) and not _norm(paragraphs[j].text).lower().startswith("contents"):
            j += 1

        if j < len(paragraphs) and _norm(paragraphs[j].text).lower().startswith("contents"):
            k = j + 1
            while k < len(paragraphs):
                pk = paragraphs[k]
                tk = _norm(pk.text)
                if not tk:
                    k += 1
                    continue
                if LO_HEAD_RE.match(tk):
                    break
                if tk.lower().startswith("assessment criteria"):
                    break
                style = (getattr(pk.style, "name", "") or "").lower()
                if style.startswith("list"):
                    lo["contents"].append({"index": len(lo["contents"]) + 1, "title": tk})
                    k += 1
                    continue
                break

        los.append(lo)
        i = j + 1
    return los


def _assemble_ia(payload_path: Path, output_path: Path) -> None:
    python_exe = Path(".venv") / "Scripts" / "python.exe"
    cmd = [str(python_exe), str(Path("tools") / "assemble_ia.py"), str(payload_path), str(output_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError((result.stdout or "") + (result.stderr or ""))


def _merge_docx(inputs: list[Path], output_path: Path) -> None:
    python_exe = Path(".venv") / "Scripts" / "python.exe"
    cmd = [str(python_exe), str(Path("tools") / "merge_docx.py"), *[str(p) for p in inputs], str(output_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError((result.stdout or "") + (result.stderr or ""))


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        print("Usage: python tools/repair_ia_course_from_processed.py <COURSE_CODE>", file=sys.stderr)
        return 2

    course_code = argv[1]
    processed_dir = PROCESSED_CBLM_DIR / course_code
    if not processed_dir.exists():
        print(f"Processed dir not found: {processed_dir}", file=sys.stderr)
        return 2

    state_dir = STATE_DIR / course_code
    out_dir = OUTPUT_DIR / course_code
    state_dir.mkdir(parents=True, exist_ok=True)
    out_dir.mkdir(parents=True, exist_ok=True)

    unit_files = sorted(processed_dir.glob("CBLM_Unit_*.docx"), key=lambda p: p.name.lower())
    if not unit_files:
        print(f"No CBLM units found in: {processed_dir}", file=sys.stderr)
        return 2

    unit_outputs: list[Path] = []
    for input_path in unit_files:
        m = re.search(r"^CBLM_Unit_(\d+)_" + re.escape(course_code) + r"_(.+)\.docx$", input_path.name, re.I)
        if not m:
            continue
        unit_index = int(m.group(1))
        title = m.group(2)

        doc = Document(str(input_path))
        head_paras = [_norm(p.text) for p in doc.paragraphs[:250] if _norm(p.text)]
        fm = _extract_front_matter_from_tables(doc)

        module_title = fm.get("module_title") or ""
        unit_of_competency = fm.get("unit_of_competency") or ""
        qual_from_tables = fm.get("qualification_title") or ""

        if not module_title or module_title.lower() in ("code", "module title"):
            module_title = _extract_module_title(head_paras)
        if not unit_of_competency or unit_of_competency.lower() in ("module title", "unit of competency", "code"):
            unit_of_competency = _extract_unit_of_competency(head_paras)

        # Qualification title: prefer table value; fallback to filename title (more reliable than paragraph heuristics).
        qual = qual_from_tables or title.replace("_", " ").strip()
        if qual.lower() in (
            "competency-based learning materials",
            "competency based learning materials",
            "how to use this competency- based learning materials",
            "how to use this competency-based learning materials",
        ):
            qual = title.replace("_", " ").strip()

        payload_path = state_dir / f"IA_Unit_{unit_index}_{_safe_filename(title)}.json"
        if payload_path.exists():
            payload = json.loads(payload_path.read_text(encoding="utf-8"))
        else:
            payload = {"qualification_title": qual, "current_unit": {}}

        payload["qualification_title"] = qual
        payload.setdefault("current_unit", {})
        payload["current_unit"]["unit_of_competency"] = unit_of_competency
        payload["current_unit"]["module_title"] = module_title
        payload["current_unit"]["next_unit_of_competency"] = _extract_next_unit(head_paras)
        payload["current_unit"]["learning_outcomes"] = _extract_learning_outcomes(doc)

        payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

        out_path = out_dir / f"IA_Unit_{unit_index}_{_safe_filename(title)}.docx"
        _assemble_ia(payload_path, out_path)
        unit_outputs.append(out_path)

    tmp_full = out_dir / "IA_FULL__tmp.docx"
    full_path = out_dir / "IA_FULL.docx"
    _merge_docx(unit_outputs, tmp_full)
    try:
        if full_path.exists():
            full_path.unlink()
        shutil.move(str(tmp_full), str(full_path))
    except PermissionError as e:
        raise RuntimeError(f"Cannot overwrite {full_path} (file may be open). Close it and rerun. Details: {e}")
    finally:
        try:
            if tmp_full.exists():
                tmp_full.unlink()
        except Exception:
            pass

    print(json.dumps({"course_code": course_code, "repaired_units": [p.name for p in unit_outputs], "full": str(full_path)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
