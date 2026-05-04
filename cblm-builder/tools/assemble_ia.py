import json
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt
from docxcompose.composer import Composer


FONT_NAME = "Bookman Old Style"
FONT_SIZE = 12

IA_TEMPLATE_FILES = {
    "front_page": "00_FrontPage.docx",
    "ia_plan": "01_iaplan.docx",
    "specific_instructions": "02_iaspesificinstruction.docx",
    "rating_sheet": "03_ratingsheet.docx",
    "questions_with_answers": "04_iaquestionwithanswer.docx",
    "recording_sheet": "05_iarecordingsheet.docx",
}


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    return str(value)


def set_run_font(run):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)


def enforce_document_font(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = FONT_NAME
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                style.font.size = Pt(FONT_SIZE)
            except Exception:
                pass


def clear_paragraph(paragraph):
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)


def write_text_to_paragraph(paragraph, text):
    style_source_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph(paragraph)
    lines = safe_text(text).split("\n")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run)
        if style_source_run is not None:
            run.bold = style_source_run.bold
            run.italic = style_source_run.italic
            run.underline = style_source_run.underline
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def replace_in_paragraphs(paragraphs, placeholder, value):
    replacement = safe_text(value)
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            write_text_to_paragraph(paragraph, paragraph.text.replace(placeholder, replacement))


def replace_everywhere(doc_part, placeholder, value):
    replace_in_paragraphs(doc_part.paragraphs, placeholder, value)
    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_everywhere(cell, placeholder, value)


def apply_scalar_map(doc, values):
    for key, value in values.items():
        replace_everywhere(doc, "{{" + key + "}}", value)


def delete_row(table, row_index):
    tbl = table._tbl
    tr = table.rows[row_index]._tr
    tbl.remove(tr)


def normalize_all_runs_font(doc_part):
    for paragraph in doc_part.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_all_runs_font(cell)


_COMMON_VERBS = {
    "apply",
    "analyze",
    "assess",
    "calculate",
    "communicate",
    "conduct",
    "configure",
    "coordinate",
    "create",
    "demonstrate",
    "describe",
    "design",
    "develop",
    "evaluate",
    "explain",
    "identify",
    "implement",
    "inspect",
    "install",
    "interpret",
    "maintain",
    "manage",
    "monitor",
    "operate",
    "perform",
    "plan",
    "prepare",
    "present",
    "record",
    "report",
    "test",
    "troubleshoot",
    "use",
}


def to_ability_phrase(content_title):
    """
    Adjust a content title so it reads correctly after:
    'The evidence must show that the trainee can ...'
    Keeps the meaning; prefers minimal changes.
    """
    title = safe_text(content_title).strip()
    if not title:
        return ""
    lower = title.lower()
    if lower.startswith("to "):
        title = title[3:].lstrip()
        lower = title.lower()
    first_word = lower.split()[0] if lower.split() else ""
    if first_word in _COMMON_VERBS:
        return title[:1].lower() + title[1:]
    # Fallback: avoid ungrammatical noun phrases after "can".
    return f"demonstrate competence in {title}"


def _get_unit(payload):
    if "current_unit" not in payload:
        raise ValueError("Payload missing required field: current_unit")
    return payload["current_unit"]


def render_front_page(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)
    apply_scalar_map(
        doc,
        {
            "unit_of_competency": unit.get("unit_of_competency", ""),
            # allow either spelling if present
            "unit_of_competency_x": unit.get("unit_of_competency", ""),
        },
    )
    normalize_all_runs_font(doc)
    return doc


def render_specific_instructions(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)
    ia = unit.get("ia", {}) or {}
    apply_scalar_map(
        doc,
        {
            "unit_of_competency": unit.get("unit_of_competency", ""),
            "module_title": unit.get("module_title", ""),
            "name_of_project": ia.get("name_of_project", ""),
            "Specific_Instructions": ia.get("Specific_Instructions", ia.get("specific_instructions", "")),
        },
    )
    normalize_all_runs_font(doc)
    return doc


def _fill_plan_or_rating_table(table, unit, adjust_contents=True, include_contents_list_col=False):
    # Identify dynamic rows (those with placeholders).
    def row_has_any(row, needles):
        return any(any(n in cell.text for n in needles) for cell in row.cells)

    # Only rows that participate in the LO/contents pattern.
    placeholders = ["{{LO}}", "{{Contents}}", "{{Contents_Z}}", "{{Y}}", "{{Z}}"]
    dynamic_row_indexes = [i for i, r in enumerate(table.rows) if row_has_any(r, placeholders)]
    # Exclude header scalar rows (e.g., unit/module) that might contain braces but aren't part of the LO/content grid.
    dynamic_row_indexes = [i for i in dynamic_row_indexes if "{{unit_of_competency}}" not in "|".join(c.text for c in table.rows[i].cells) and "{{module_title}}" not in "|".join(c.text for c in table.rows[i].cells)]

    los = unit.get("learning_outcomes", []) or []
    lo_iter = iter(los)
    current_lo = None
    current_contents_iter = iter(())

    used_row_indexes = []

    for row_index in dynamic_row_indexes:
        row = table.rows[row_index]
        row_text = "|".join(cell.text for cell in row.cells)

        if "{{LO}}" in row_text:
            current_lo = next(lo_iter, None)
            current_contents_iter = iter((current_lo.get("contents", []) or []) if current_lo else ())
            if not current_lo:
                break

            y = current_lo.get("index", "")
            lo_title = current_lo.get("title", "")
            for cell in row.cells:
                text = cell.text.replace("{{Y}}", safe_text(y)).replace("{{LO}}", safe_text(lo_title))
                write_text_to_paragraph(cell.paragraphs[0], text)
                for p in cell.paragraphs[1:]:
                    p._element.getparent().remove(p._element)
            used_row_indexes.append(row_index)
            continue

        # Content row: requires an active LO.
        if not current_lo:
            break

        content = next(current_contents_iter, None)
        if not content:
            # No more contents for this LO; stop consuming content rows until the next LO row.
            # Leave row for later deletion.
            continue

        y = current_lo.get("index", "")
        z = content.get("index", "")
        content_title = safe_text(content.get("title", ""))
        content_for_contents = to_ability_phrase(content_title) if adjust_contents else content_title

        contents_list = "\n".join(
            safe_text(c.get("title", "")).strip()
            for c in (current_lo.get("contents", []) or [])
            if safe_text(c.get("title", "")).strip()
        )

        for cell in row.cells:
            text = cell.text.replace("{{Y}}", safe_text(y)).replace("{{Z}}", safe_text(z))
            text = text.replace("{{Contents}}", content_for_contents)
            if include_contents_list_col:
                text = text.replace("{{Contents_Z}}", contents_list)
            write_text_to_paragraph(cell.paragraphs[0], text)
            for p in cell.paragraphs[1:]:
                p._element.getparent().remove(p._element)

        used_row_indexes.append(row_index)

    # Delete all placeholder rows that weren't used.
    used = set(used_row_indexes)
    for row_index in reversed(dynamic_row_indexes):
        if row_index not in used:
            delete_row(table, row_index)


def render_ia_plan(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)

    apply_scalar_map(
        doc,
        {
            "unit_of_competency": unit.get("unit_of_competency", ""),
            "module_title": unit.get("module_title", ""),
        },
    )

    # Table 0 contains LO/content listing.
    if doc.tables:
        table = doc.tables[0]
        _fill_plan_or_rating_table(
            table,
            unit,
            adjust_contents=True,  # ability phrasing for {{Contents}}
            include_contents_list_col=True,  # fill {{Contents_Z}} with raw list of titles
        )

    normalize_all_runs_font(doc)
    return doc


def render_rating_sheet(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)
    ia = unit.get("ia", {}) or {}
    apply_scalar_map(
        doc,
        {
            "qualification_title": payload.get("qualification_title", ""),
            "qualification": payload.get("qualification_title", ""),
            "module_title": unit.get("module_title", ""),
            "instructions_for_demo": ia.get("instructions_for_demo", ""),
            "list_of_materials_and_equipment": ia.get("list_of_materials_and_equipment", ""),
        },
    )

    if doc.tables:
        table = doc.tables[0]
        _fill_plan_or_rating_table(
            table,
            unit,
            adjust_contents=True,  # keep consistent: content phrasing for {{Contents}}
            include_contents_list_col=False,
        )

    normalize_all_runs_font(doc)
    return doc


def render_questions_with_answers(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)
    ia = unit.get("ia", {}) or {}
    oral = ia.get("oral_questions", []) or []

    values = {}
    for i in range(1, 6):
        q = oral[i - 1] if i - 1 < len(oral) else {}
        values[f"oral_question_{i}"] = q.get("question", "")
        values[f"oral_question_{i}_acceptable_answer"] = q.get("acceptable_answer", q.get("answer", ""))

    apply_scalar_map(doc, values)
    normalize_all_runs_font(doc)
    return doc


def render_recording_sheet(template_path, payload):
    doc = Document(template_path)
    enforce_document_font(doc)
    unit = _get_unit(payload)
    ia = unit.get("ia", {}) or {}
    interview = ia.get("interview_questions", []) or []

    values = {
        "module_title": unit.get("module_title", ""),
        "qualification": payload.get("qualification_title", ""),
        "qualification_title": payload.get("qualification_title", ""),
    }
    for i in range(1, 6):
        values[f"interview_question_{i}"] = interview[i - 1] if i - 1 < len(interview) else ""

    apply_scalar_map(doc, values)
    normalize_all_runs_font(doc)
    return doc


def assemble_ia(payload_path, output_path, templates_dir="templates/IA TEMPLATES"):
    payload_path = Path(payload_path)
    output_path = Path(output_path)
    templates_dir = Path(templates_dir)

    payload = json.loads(payload_path.read_text(encoding="utf-8"))

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)

        parts = []
        parts.append(render_front_page(templates_dir / IA_TEMPLATE_FILES["front_page"], payload))
        parts.append(render_ia_plan(templates_dir / IA_TEMPLATE_FILES["ia_plan"], payload))
        parts.append(render_specific_instructions(templates_dir / IA_TEMPLATE_FILES["specific_instructions"], payload))
        parts.append(render_rating_sheet(templates_dir / IA_TEMPLATE_FILES["rating_sheet"], payload))
        parts.append(render_questions_with_answers(templates_dir / IA_TEMPLATE_FILES["questions_with_answers"], payload))
        parts.append(render_recording_sheet(templates_dir / IA_TEMPLATE_FILES["recording_sheet"], payload))

        base_path = temp_dir / "000_ia_base.docx"
        parts[0].save(base_path)
        base_doc = Document(base_path)
        composer = Composer(base_doc)
        for part in parts[1:]:
            # Ensure each template starts on a new page.
            breaker = composer.doc.add_paragraph().add_run()
            breaker.add_break(WD_BREAK.PAGE)
            part_path = temp_dir / f"part_{len(composer.doc.paragraphs)}.docx"
            part.save(part_path)
            composer.append(Document(part_path))

        output_path.parent.mkdir(parents=True, exist_ok=True)
        composer.save(str(output_path))


def main(argv):
    if len(argv) < 3:
        print("Usage: python tools/assemble_ia.py <payload.json> <output.docx>", file=sys.stderr)
        return 2
    assemble_ia(argv[1], argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
