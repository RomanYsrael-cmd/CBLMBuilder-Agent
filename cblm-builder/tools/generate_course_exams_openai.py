import argparse
import json
import os
import random
import re
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt
from docxcompose.composer import Composer


FONT_NAME = "Arial Narrow"
FONT_SIZE = 12


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(v) for v in value)
    return str(value)


def set_run_font(run):
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)


def enforce_document_font(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = FONT_NAME
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                style.font.size = Pt(FONT_SIZE)
            except Exception:
                pass


def clear_paragraph(paragraph):
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)


def write_text_to_paragraph(paragraph, text):
    style_source_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph(paragraph)
    lines = safe_text(text).split("\n")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run)
        if style_source_run is not None:
            run.bold = style_source_run.bold
            run.italic = style_source_run.italic
            run.underline = style_source_run.underline
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def replace_in_paragraphs(paragraphs, placeholder, value):
    replacement = safe_text(value)
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            write_text_to_paragraph(paragraph, paragraph.text.replace(placeholder, replacement))


def replace_everywhere(doc_part, placeholder, value):
    replace_in_paragraphs(doc_part.paragraphs, placeholder, value)
    for table in getattr(doc_part, "tables", []) or []:
        for row in table.rows:
            for cell in row.cells:
                replace_everywhere(cell, placeholder, value)


def apply_scalar_map(doc_part, values):
    for key, value in values.items():
        replace_everywhere(doc_part, "{{" + key + "}}", value)


def normalize_all_runs_font(doc_part):
    for paragraph in doc_part.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in getattr(doc_part, "tables", []) or []:
        for row in table.rows:
            for cell in row.cells:
                normalize_all_runs_font(cell)


def load_payload(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _infer_course_title(payloads: list[dict]) -> str:
    for p in payloads:
        t = safe_text(p.get("qualification_title", "")).strip()
        if t:
            return t
    return ""


def _term_unit_ranges(unit_count: int, terms: list[str]) -> list[tuple[str, int, int]]:
    term_count = len(terms)
    base = unit_count // term_count
    remainder = unit_count % term_count
    ranges = []
    start = 0
    for i, term in enumerate(terms):
        size = base + (1 if i < remainder else 0)
        end = start + size
        ranges.append((term, start, end))
        start = end
    return ranges


def _outline_for_payload(payload: dict, *, max_keyfacts_words: int = 220) -> str:
    """
    Build a compact outline to condition the model without sending full Key Facts.
    """
    unit = payload.get("current_unit") or {}
    parts = []
    parts.append(f"Unit of Competency: {safe_text(unit.get('unit_of_competency', '')).strip()}")
    parts.append(f"Module Title: {safe_text(unit.get('module_title', '')).strip()}")
    for lo in unit.get("learning_outcomes") or []:
        lo_title = safe_text(lo.get("title", "")).strip()
        if not lo_title:
            continue
        parts.append(f"- Topic: {lo_title}")
        subtopics = [safe_text(c.get("title", "")).strip() for c in (lo.get("contents") or []) if safe_text(c.get("title", "")).strip()]
        if subtopics:
            parts.append("  Subtopics: " + "; ".join(subtopics))
        kf = safe_text(lo.get("key_facts", "")).strip()
        if kf:
            words = kf.split()
            excerpt = " ".join(words[:max_keyfacts_words])
            parts.append("  Key Facts excerpt: " + excerpt)
    return "\n".join(parts).strip()


def _openai_responses_json_schema(*, api_key: str, model: str, prompt: str, schema: dict, timeout_s: int = 600) -> dict:
    url = "https://api.openai.com/v1/responses"
    data = {
        "model": model,
        "input": [
            {
                "role": "system",
                "content": [
                    {
                        "type": "input_text",
                        "text": "You are an exam writer for competency-based training. Output must be JSON that matches the given schema.",
                    }
                ],
            },
            {"role": "user", "content": [{"type": "input_text", "text": prompt}]},
        ],
        "text": {"format": {"type": "json_schema", "name": "exam_mcqs", "schema": schema, "strict": True}},
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(data).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout_s) as resp:
            raw = resp.read().decode("utf-8")
    except urllib.error.HTTPError as e:
        body = ""
        try:
            body = e.read().decode("utf-8", errors="replace")
        except Exception:
            body = str(e)
        raise RuntimeError(f"OpenAI API error {e.code}: {body[:1000]}") from e

    payload = json.loads(raw)
    # The structured output JSON will be in output_text.
    out_text = payload.get("output_text", "")
    if not out_text:
        raise RuntimeError("OpenAI API returned empty output_text.")
    try:
        return json.loads(out_text)
    except Exception as e:
        raise RuntimeError(f"Failed to parse JSON output_text: {e}. Raw: {out_text[:1000]}") from e


def generate_mcqs_for_term(
    *,
    api_key: str,
    model: str,
    course_code: str,
    course_title: str,
    term: str,
    scoped_outlines: list[str],
    seed: int,
) -> list[dict]:
    schema = {
        "type": "object",
        "properties": {
            "mcqs": {
                "type": "array",
                "minItems": 50,
                "maxItems": 50,
                "items": {
                    "type": "object",
                    "properties": {
                        "question": {"type": "string", "minLength": 1},
                        "a": {"type": "string", "minLength": 1},
                        "b": {"type": "string", "minLength": 1},
                        "c": {"type": "string", "minLength": 1},
                        "d": {"type": "string", "minLength": 1},
                        "answer": {"type": "string", "enum": ["A", "B", "C", "D"]},
                    },
                    "required": ["question", "a", "b", "c", "d", "answer"],
                    "additionalProperties": False,
                },
            }
        },
        "required": ["mcqs"],
        "additionalProperties": False,
    }

    context = "\n\n".join(scoped_outlines).strip()
    prompt = "\n".join(
        [
            f"Write a {term} exam for:",
            f"Course Code: {course_code}",
            f"Course Title: {course_title}",
            "",
            "Requirements:",
            "- Output EXACTLY 50 multiple-choice questions.",
            "- Each question has 4 choices (A-D) and exactly one correct answer.",
            "- Questions must be based ONLY on the provided unit/topic outlines and Key Facts excerpts.",
            "- Create a DIFFERENT set from any Let’s Exercise questions (do not copy or paraphrase them).",
            "- Mix difficulty: 15 easy, 20 moderate, 15 challenging (do not label difficulty).",
            "- Avoid repetitive stems and avoid duplicates.",
            "",
            f"Randomization seed: {seed}",
            "",
            "Unit/Topic outlines:",
            context,
        ]
    ).strip()

    result = _openai_responses_json_schema(api_key=api_key, model=model, prompt=prompt, schema=schema)
    mcqs = result.get("mcqs") or []
    if len(mcqs) != 50:
        raise RuntimeError(f"Model returned {len(mcqs)} MCQs, expected 50.")
    return mcqs


def generate_midterm_and_finals(
    *,
    api_key: str,
    model: str,
    course_code: str,
    course_title: str,
    outlines_by_unit: list[str],
    seed: int,
) -> dict:
    """
    Returns:
      {"course_code","course_title","midterm_mcqs","finals_mcqs"}
    Term scoping splits units in order (first half -> MIDTERM, second half -> FINAL; remainder to MIDTERM).
    """
    terms = ["MIDTERM", "FINAL"]
    ranges = _term_unit_ranges(len(outlines_by_unit), terms)
    midterm_mcqs: list[dict] = []
    finals_mcqs: list[dict] = []
    for term, start, end in ranges:
        scoped = outlines_by_unit[start:end]
        if not scoped:
            raise RuntimeError(f"Term '{term}' has empty unit scope; check unit count.")
        term_seed = int(seed or 0) + (1 if term == "MIDTERM" else 2)
        mcqs = generate_mcqs_for_term(
            api_key=api_key,
            model=model,
            course_code=course_code,
            course_title=course_title,
            term=term,
            scoped_outlines=scoped,
            seed=term_seed,
        )
        if term == "MIDTERM":
            midterm_mcqs = mcqs
        else:
            finals_mcqs = mcqs
    return {"course_code": course_code, "course_title": course_title, "midterm_mcqs": midterm_mcqs, "finals_mcqs": finals_mcqs}


def render_exam_docx(template_path: Path, *, course_code: str, course_title: str, mcqs: list[dict]) -> Document:
    doc = Document(str(template_path))
    enforce_document_font(doc)
    apply_scalar_map(doc, {"COURSE_CODE": course_code, "COURSE_TITLE": course_title, "TERM": ""})

    # Locate question table (expected 25 rows x 2 columns; total 50 cells).
    question_table = None
    for table in doc.tables:
        if len(table.rows) == 25 and len(table.columns) == 2:
            question_table = table
            break
    if question_table is None:
        raise ValueError("Could not locate question table (expected 25x2).")

    cells = []
    for row in question_table.rows:
        cells.append(row.cells[0])
        cells.append(row.cells[1])

    if len(mcqs) != 50:
        raise ValueError(f"Expected exactly 50 MCQs, got {len(mcqs)}")

    for idx, item in enumerate(mcqs, start=1):
        apply_scalar_map(
            cells[idx - 1],
            {
                "Q_NUM": str(idx),
                "QUESTION": safe_text(item.get("question", "")).strip(),
                "Q_CHOICE1": safe_text(item.get("a", "")).strip(),
                "Q_CHOICE2": safe_text(item.get("b", "")).strip(),
                "Q_CHOICE3": safe_text(item.get("c", "")).strip(),
                "Q_CHOICE4": safe_text(item.get("d", "")).strip(),
            },
        )

    normalize_all_runs_font(doc)
    return doc


def append_docx(base_path: Path, append_paths: list[Path], out_path: Path) -> None:
    base_doc = Document(str(base_path))
    composer = Composer(base_doc)
    for p in append_paths:
        composer.doc.add_page_break()
        composer.append(Document(str(p)))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate MIDTERM and FINAL exam DOCX using IA templates (03_midterm.docx / 04_finals.docx).")
    parser.add_argument("--course-code", required=True)
    parser.add_argument("--course-title", default="", help="Defaults to qualification_title from payloads if omitted.")
    parser.add_argument("--payloads", type=Path, nargs="+", required=True, help="All unit payload JSON files (UC order).")
    parser.add_argument("--templates-dir", type=Path, default=Path("templates/IA TEMPLATES"))
    parser.add_argument("--outdir", type=Path, default=Path("output/ia"))
    parser.add_argument("--state-dir", type=Path, default=Path("state/exams"))
    parser.add_argument("--model", default=os.environ.get("OPENAI_EXAM_MODEL") or "gpt-4o-mini")
    parser.add_argument("--seed", type=int, default=0)
    parser.add_argument("--sleep-s", type=float, default=0.0)
    parser.add_argument("--append-to-ia-full", action="store_true", help="Append exams to output/ia/<COURSE_CODE>/IA_FULL.docx if present.")
    args = parser.parse_args()

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("Missing OPENAI_API_KEY environment variable.", file=sys.stderr)
        return 2

    payload_objs = [load_payload(p) for p in args.payloads]
    course_title = args.course_title.strip() or _infer_course_title(payload_objs) or args.course_code

    terms = ["MIDTERM", "FINAL"]
    ranges = _term_unit_ranges(len(payload_objs), terms)

    templates_dir: Path = args.templates_dir
    midterm_template = templates_dir / "03_midterm.docx"
    finals_template = templates_dir / "04_finals.docx"

    course_outdir = args.outdir / args.course_code
    course_outdir.mkdir(parents=True, exist_ok=True)
    course_statedir = args.state_dir / args.course_code
    course_statedir.mkdir(parents=True, exist_ok=True)

    generated_paths: list[Path] = []
    exam_payload = {"course_code": args.course_code, "course_title": course_title, "midterm_mcqs": [], "finals_mcqs": []}

    for term, start, end in ranges:
        scoped = payload_objs[start:end]
        outlines = [_outline_for_payload(p) for p in scoped if p]
        seed = int(args.seed or 0) + (1 if term == "MIDTERM" else 2)
        mcqs = generate_mcqs_for_term(
            api_key=api_key,
            model=args.model,
            course_code=args.course_code,
            course_title=course_title,
            term=term,
            scoped_outlines=outlines,
            seed=seed,
        )

        json_path = course_statedir / f"{term}.json"
        json_path.write_text(json.dumps({"term": term, "course_code": args.course_code, "course_title": course_title, "mcqs": mcqs}, ensure_ascii=False, indent=2), encoding="utf-8")

        if term == "MIDTERM":
            exam_payload["midterm_mcqs"] = mcqs
            doc = render_exam_docx(midterm_template, course_code=args.course_code, course_title=course_title, mcqs=mcqs)
            out_path = course_outdir / "03_midterm.docx"
        else:
            exam_payload["finals_mcqs"] = mcqs
            doc = render_exam_docx(finals_template, course_code=args.course_code, course_title=course_title, mcqs=mcqs)
            out_path = course_outdir / "04_finals.docx"

        doc.save(str(out_path))
        generated_paths.append(out_path)

        if args.sleep_s > 0:
            time.sleep(args.sleep_s)

    (course_statedir / "EXAMS_PAYLOAD.json").write_text(json.dumps(exam_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    if args.append_to_ia_full:
        ia_full = course_outdir / "IA_FULL.docx"
        if ia_full.exists():
            tmp = course_outdir / "IA_FULL__with_exams__tmp.docx"
            append_docx(ia_full, generated_paths, tmp)
            final = course_outdir / "IA_FULL.docx"
            try:
                if final.exists():
                    final.unlink()
                tmp.rename(final)
            except PermissionError as e:
                raise RuntimeError(f"Cannot overwrite {final} (file may be open). Close it and rerun. Details: {e}")
        else:
            print(f"Note: IA_FULL.docx not found at {ia_full}; generated exams only.", file=sys.stderr)

    print(json.dumps({"course_code": args.course_code, "course_title": course_title, "outputs": [str(p) for p in generated_paths]}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
