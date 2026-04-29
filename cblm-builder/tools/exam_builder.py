import argparse
import json
import random
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAME = "Arial Narrow"
FONT_SIZE = 9


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(v) for v in value)
    return str(value)


def normalize(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9\\s]", " ", text)
    text = re.sub(r"\\s+", " ", text).strip()
    return text


def shingles(tokens: list[str], size: int = 5) -> set[str]:
    if not tokens:
        return set()
    if len(tokens) < size:
        return {" ".join(tokens)}
    return {" ".join(tokens[i : i + size]) for i in range(len(tokens) - size + 1)}


def jaccard(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    return len(left & right) / len(left | right)


def tokenize(text: str) -> list[str]:
    return [t for t in normalize(text).split(" ") if t]


@dataclass
class MCQ:
    question: str
    a: str
    b: str
    c: str
    d: str
    answer: str | None = None
    source: str | None = None
    level: str | None = None


def parse_exercise_questions(value) -> list[MCQ]:
    """
    Accepts either:
    - list of dicts: [{"question":..,"choices":[..], "answer":..}, ...]
    - list of strings in a simple format
    - a single string containing multiple questions
    """
    if value is None:
        return []

    if isinstance(value, list):
        if value and isinstance(value[0], dict):
            out = []
            for item in value:
                q = safe_text(item.get("question")).strip()
                choices = item.get("choices") or []
                if len(choices) != 4:
                    continue
                out.append(
                    MCQ(
                        question=q,
                        a=safe_text(choices[0]).strip(),
                        b=safe_text(choices[1]).strip(),
                        c=safe_text(choices[2]).strip(),
                        d=safe_text(choices[3]).strip(),
                        answer=safe_text(item.get("answer")).strip() or None,
                        level=safe_text(item.get("level")).strip().lower() or None,
                    )
                )
            return out

        if value and isinstance(value[0], str):
            blob = "\n".join(value)
            return parse_exercise_questions(blob)

        return []

    blob = safe_text(value)
    if not blob.strip():
        return []

    # Very permissive parser for:
    # 1. Question text?
    # A. ...
    # B. ...
    # C. ...
    # D. ...
    # Answer: X
    lines = [ln.rstrip() for ln in blob.splitlines() if ln.strip()]
    out: list[MCQ] = []
    i = 0
    while i < len(lines):
        q_line = lines[i]
        # Skip numbering like "1." or "1)"
        q_line = re.sub(r"^\\s*\\d+\\s*[\\.)]\\s*", "", q_line).strip()
        # Find next 4 choice lines
        if i + 4 >= len(lines):
            break
        a = re.sub(r"^\\s*A\\.?\\s*", "", lines[i + 1]).strip()
        b = re.sub(r"^\\s*B\\.?\\s*", "", lines[i + 2]).strip()
        c = re.sub(r"^\\s*C\\.?\\s*", "", lines[i + 3]).strip()
        d = re.sub(r"^\\s*D\\.?\\s*", "", lines[i + 4]).strip()
        answer = None
        j = i + 5
        if j < len(lines) and re.match(r"^\\s*answer\\s*[:\\-]", lines[j], re.IGNORECASE):
            answer = re.sub(r"^\\s*answer\\s*[:\\-]\\s*", "", lines[j], flags=re.IGNORECASE).strip()
            j += 1
        out.append(MCQ(question=q_line, a=a, b=b, c=c, d=d, answer=answer))
        i = j

    return out


def load_payload(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def content_title_tokens(payload: dict) -> set[str]:
    tokens: set[str] = set()
    unit = payload.get("current_unit") or {}
    for lo in unit.get("learning_outcomes") or []:
        for content in lo.get("contents") or []:
            title = safe_text(content.get("title")).strip()
            for t in tokenize(title):
                if len(t) >= 5:
                    tokens.add(t)
    return tokens


def extract_mcqs_from_payload(payload: dict, source: str) -> list[MCQ]:
    unit = payload.get("current_unit") or {}
    out: list[MCQ] = []
    for lo in unit.get("learning_outcomes") or []:
        for content in lo.get("contents") or []:
            value = content.get("exercise_questions")
            mcqs = parse_exercise_questions(value)
            for m in mcqs:
                m.source = source
            out.extend(mcqs)
    return out


def validate_mcqs(mcqs: list[MCQ], required_tokens: set[str], *, strict_relevance: bool = True) -> None:
    if not mcqs:
        raise ValueError("No MCQs found.")

    norm_questions = [normalize(m.question) for m in mcqs]
    if any(not q for q in norm_questions):
        raise ValueError("Found empty question text.")

    # Ensure no duplicates and not too-similar questions.
    seen = set()
    for q in norm_questions:
        if q in seen:
            raise ValueError("Duplicate questions detected.")
        seen.add(q)

    # Disallow internal tags / unit-scoped meta phrasing; exams should read naturally.
    for idx, m in enumerate(mcqs, start=1):
        raw = (m.question or "").strip()
        raw_l = raw.lower()
        if re.search(r"\[[^\]]+\]", raw):
            raise ValueError(f"Question {idx} contains bracketed metadata; remove internal tags.")
        if re.search(r"\\buc\\s*\\d+\\b", raw_l) or "unit of competency" in raw_l:
            raise ValueError(f"Question {idx} contains UC/unit meta text; remove UC references.")
        if any(term in raw_l for term in [" learning outcome ", " subtopic ", " topic "]):
            raise ValueError(f"Question {idx} contains structural syllabus metadata; rewrite as a natural exam question.")

    prepared = [shingles(tokenize(m.question)) for m in mcqs]
    for i in range(len(mcqs)):
        for j in range(i + 1, len(mcqs)):
            if jaccard(prepared[i], prepared[j]) >= 0.4:
                raise ValueError(f"Questions too similar: {i+1} vs {j+1}")

    # Scaffold-level repetition checks live in validate_exam_docx.py; here we only enforce
    # hard duplication and pairwise similarity on the full question text.

    # Basic choice validation
    for idx, m in enumerate(mcqs, start=1):
        choices = [m.a, m.b, m.c, m.d]
        if any(not c.strip() for c in choices):
            raise ValueError(f"Blank choice in question {idx}")
        if len({normalize(c) for c in choices}) != 4:
            raise ValueError(f"Repeated choices in question {idx}")

    if strict_relevance and required_tokens:
        missing = []
        for idx, m in enumerate(mcqs, start=1):
            q_tokens = set(tokenize(m.question))
            if not (q_tokens & required_tokens):
                missing.append(idx)
        if missing:
            raise ValueError(f"Relevance check failed; questions missing content-keyword overlap: {missing[:10]}")


def replace_placeholders_in_paragraph(paragraph, mapping: dict[str, str]):
    if not paragraph.text:
        return
    text = paragraph.text
    touched = False
    for k, v in mapping.items():
        ph = "{{" + k + "}}"
        if ph in text:
            text = text.replace(ph, v)
            touched = True
    if not touched:
        return

    # rewrite with single run, preserve line breaks
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        if i < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def replace_placeholders_in_cell(cell, mapping: dict[str, str]):
    for paragraph in cell.paragraphs:
        replace_placeholders_in_paragraph(paragraph, mapping)


def set_run_font(run):
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)


def enforce_document_font(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = FONT_NAME
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                style.font.size = Pt(FONT_SIZE)
            except Exception:
                pass


def normalize_all_runs_font(doc_part):
    for paragraph in doc_part.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in doc_part.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_all_runs_font(cell)


def fill_exam_template(template_path: Path, output_path: Path, *, course_code: str, course_title: str, term: str, mcqs: list[MCQ]) -> None:
    doc = Document(template_path)
    enforce_document_font(doc)

    # Global placeholders for header
    header_map = {"COURSE_CODE": course_code, "COURSE_TITLE": course_title, "TERM": term}
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, header_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, header_map)

    # Question table: last table with 25x2.
    question_table = None
    for table in doc.tables:
        if len(table.rows) == 25 and len(table.columns) == 2:
            question_table = table
            break
    if question_table is None:
        raise ValueError("Could not locate question table (expected 25x2).")

    cells = []
    for row in question_table.rows:
        cells.append(row.cells[0])
        cells.append(row.cells[1])

    if len(mcqs) > len(cells):
        raise ValueError(f"Too many questions for template: {len(mcqs)} > {len(cells)}")

    # Fill questions sequentially
    for idx, mcq in enumerate(mcqs, start=1):
        mapping = {
            "Q_NUM": str(idx),
            "QUESTION": mcq.question.strip(),
            "Q_CHOICE1": mcq.a.strip(),
            "Q_CHOICE2": mcq.b.strip(),
            "Q_CHOICE3": mcq.c.strip(),
            "Q_CHOICE4": mcq.d.strip(),
        }
        replace_placeholders_in_cell(cells[idx - 1], mapping)

    # Clear any remaining placeholders for unused cells
    for idx in range(len(mcqs), len(cells)):
        mapping = {"Q_NUM": "", "QUESTION": "", "Q_CHOICE1": "", "Q_CHOICE2": "", "Q_CHOICE3": "", "Q_CHOICE4": ""}
        replace_placeholders_in_cell(cells[idx], mapping)

    normalize_all_runs_font(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build exam DOCX files from unit payload JSONs and the EXAM TEMPLATE.docx.")
    parser.add_argument("--course-code", required=True)
    parser.add_argument("--course-title", required=True)
    parser.add_argument("--template", type=Path, default=Path("templates/EXAM TEMPLATE.docx"))
    parser.add_argument("--outdir", type=Path, default=Path("output/exams"))
    parser.add_argument("--terms", default="MIDTERM,FINAL", help="Comma-separated terms to generate (e.g., PRELIMINARY,MIDTERM,PRE-FINAL,FINAL)")
    parser.add_argument("--payloads", type=Path, nargs="+", required=True, help="Payload JSON files for all units (in UC order).")
    parser.add_argument("--questions-per-exam", type=int, default=50)
    parser.add_argument("--seed", type=int, default=0)
    parser.add_argument(
        "--taxonomy",
        default="10,10,30",
        help="Comma-separated counts for knowledge,comprehension,application (default 10,10,30).",
    )
    args = parser.parse_args()

    terms = [t.strip() for t in args.terms.split(",") if t.strip()]
    if not terms:
        raise ValueError("No terms specified.")

    payloads = [load_payload(p) for p in args.payloads]
    if not payloads:
        raise ValueError("No payloads provided.")

    # Divide units among terms in order.
    unit_count = len(payloads)
    term_count = len(terms)
    base = unit_count // term_count
    remainder = unit_count % term_count
    ranges = []
    start = 0
    for i in range(term_count):
        size = base + (1 if i < remainder else 0)
        end = start + size
        ranges.append((start, end))
        start = end

    rng = random.Random(args.seed or 0)
    outdir = args.outdir
    outdir.mkdir(parents=True, exist_ok=True)

    for term, (start, end) in zip(terms, ranges, strict=True):
        scoped = payloads[start:end]
        if not scoped:
            raise ValueError(f"Term '{term}' has empty unit scope; adjust terms or unit count.")

        token_union: set[str] = set()
        pool: list[MCQ] = []
        for idx, p in enumerate(scoped, start=start + 1):
            token_union |= content_title_tokens(p)
            pool.extend(extract_mcqs_from_payload(p, source=f"UC{idx}"))

        # Choose questions
        pool = [m for m in pool if m.question.strip()]
        rng.shuffle(pool)

        # Taxonomy selection: take from the pool in order, distributing by requested buckets.
        try:
            k_cnt, c_cnt, a_cnt = [int(x.strip()) for x in str(args.taxonomy).split(",")]
        except Exception as e:
            raise ValueError(f"Invalid --taxonomy value: {args.taxonomy}") from e
        if k_cnt + c_cnt + a_cnt != args.questions_per_exam:
            raise ValueError("taxonomy counts must sum to questions-per-exam")

        # If upstream items provide levels, enforce the requested mix strictly.
        level_values = {m.level for m in pool if m.level}
        if level_values:
            knowledge = [m for m in pool if m.level == "knowledge"]
            comprehension = [m for m in pool if m.level == "comprehension"]
            application = [m for m in pool if m.level == "application"]
            if len(knowledge) < k_cnt or len(comprehension) < c_cnt or len(application) < a_cnt:
                raise ValueError(
                    f"Not enough tagged questions for taxonomy mix in term '{term}': "
                    f"knowledge={len(knowledge)}, comprehension={len(comprehension)}, application={len(application)}"
                )
            rng.shuffle(knowledge)
            rng.shuffle(comprehension)
            rng.shuffle(application)
            selected = knowledge[:k_cnt] + comprehension[:c_cnt] + application[:a_cnt]
            rng.shuffle(selected)
        else:
            selected = pool[: args.questions_per_exam]

        if len(selected) < args.questions_per_exam:
            raise ValueError(f"Not enough MCQs in scope for term '{term}': {len(selected)} < {args.questions_per_exam}")

        validate_mcqs(selected, token_union, strict_relevance=True)

        safe_term = term.replace("/", "-").strip()
        output_path = outdir / f"EXAM_{args.course_code}_{safe_term}.docx"
        fill_exam_template(
            args.template,
            output_path,
            course_code=args.course_code,
            course_title=args.course_title,
            term=term,
            mcqs=selected,
        )
        print(str(output_path))

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
