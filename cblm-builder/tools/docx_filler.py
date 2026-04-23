from docx import Document


def replace_in_paragraphs(paragraphs, placeholder, value):
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)


def replace_in_table(table, placeholder, value):
    for row in table.rows:
        for cell in row.cells:
            replace_everywhere(cell, placeholder, value)


def replace_in_header_footer(part, placeholder, value):
    replace_in_paragraphs(part.paragraphs, placeholder, value)
    for table in part.tables:
        replace_in_table(table, placeholder, value)


def replace_everywhere(doc_part, placeholder, value):
    replace_in_paragraphs(doc_part.paragraphs, placeholder, value)
    for table in doc_part.tables:
        replace_in_table(table, placeholder, value)


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(v) for v in value)
    return str(value)


def fill_docx(template_path, output_path, data):
    doc = Document(template_path)

    for key, value in data.items():
        placeholder = "{{" + key + "}}"
        text_value = safe_text(value)

        # Main document body
        replace_everywhere(doc, placeholder, text_value)

        # Headers and footers in every section
        for section in doc.sections:
            replace_in_header_footer(section.header, placeholder, text_value)
            replace_in_header_footer(section.footer, placeholder, text_value)

            # First page / even page variants if present
            if hasattr(section, "first_page_header"):
                replace_in_header_footer(section.first_page_header, placeholder, text_value)
            if hasattr(section, "first_page_footer"):
                replace_in_header_footer(section.first_page_footer, placeholder, text_value)
            if hasattr(section, "even_page_header"):
                replace_in_header_footer(section.even_page_header, placeholder, text_value)
            if hasattr(section, "even_page_footer"):
                replace_in_header_footer(section.even_page_footer, placeholder, text_value)

    doc.save(output_path)
    return output_path